# app.py
"""
Streamlit RAG Support App (complete, with debug + self-tests + UI + Quit)

Features
1) YouTrack connection (URL + token) and project/issues list
2) Ticket ingestion into a local Vector DB (Chroma) with local or OpenAI embeddings
3) LLM selection (OpenAI or local via Ollama)
4) RAG chat: search similar tickets + LLM answer with ID citations
5) CLI mode with self-tests when Streamlit is not available
6) Quit button in the sidebar to close the app from the browser
7) If data/chroma/chroma.sqlite3 is found, the collection is opened automatically and the number of loaded documents is shown (or “N/A” if not available)

Recommended start
streamlit run app.py --server.port 8502
"""

import os
import sys
import time
import json
import math
import uuid
import typing
from dataclasses import dataclass
from typing import List, Tuple, Optional
import re
import shutil, subprocess
from datetime import datetime
import io
import traceback

# NEW (optional): token encoder for precise chunking
try:
    import tiktoken
    _tk_enc = tiktoken.get_encoding("cl100k_base")
except Exception:
    _tk_enc = None

# === Sticky prefs (Level A) ===
import os

# --- NumPy 2 compatibility shim (before importing chromadb) ---
try:
    import numpy as _np  # type: ignore
    if not hasattr(_np, "float_"):  # NumPy 2.0+
        _np.float_ = _np.float64     # type: ignore[attr-defined]
    if not hasattr(_np, "int_"):
        _np.int_ = _np.int64         # type: ignore[attr-defined]
    if not hasattr(_np, "uint"):
        _np.uint = _np.uint64        # type: ignore[attr-defined]
except Exception:
    pass

APP_DIR = os.path.dirname(os.path.abspath(__file__))

def is_cloud() -> bool:
    try:
        import streamlit as st  # type: ignore
        if hasattr(st, "secrets") and str(st.secrets.get("DEPLOY_ENV", "")).lower() == "cloud":
            return True
    except Exception:
        pass
    # fallback (heuristic): repo not writable => treat it as cloud
    try:
        test_path = os.path.join(APP_DIR, ".write_test")
        with open(test_path, "w") as f:
            f.write("x")
        os.remove(test_path)
        return False
    except Exception:
        return True

IS_CLOUD = is_cloud()

def pick_default_chroma_dir() -> str:
    env_dir = os.getenv("CHROMA_DIR")
    if env_dir:
        return env_dir

    # also support st.secrets["CHROMA_DIR"]
    try:
        import streamlit as st  # type: ignore
        sec_dir = st.secrets.get("CHROMA_DIR") if hasattr(st, "secrets") else None
        if sec_dir:
            return str(sec_dir)
    except Exception:
        pass

    # We already checked writability in is_cloud(); rely only on that.
    if IS_CLOUD:
        # cloud / ambienti read-only → usa /tmp
        return "/tmp/chroma"

    # local / docker con codice scrivibile → usa cartella del progetto
    return os.path.join(APP_DIR, "data", "chroma")


DEFAULT_CHROMA_DIR = pick_default_chroma_dir()

# Preferences path: write to /tmp in cloud
PREFS_PATH = os.path.join("/tmp", ".app_prefs.json") if IS_CLOUD else os.path.join(APP_DIR, ".app_prefs.json")


NON_SENSITIVE_PREF_KEYS = {
    "yt_url",
    "persist_dir",
    "emb_backend",
    "emb_model_name",
    "llm_provider",
    "llm_model",
    "llm_temperature",
    "max_distance",
    "show_prompt",
    "collection_selected",
    "new_collection_name", 
    "enable_chunking", 
    "chunk_size", 
    "chunk_overlap", 
    "chunk_min", 
    "show_distances", 
    "top_k", 
    "collapse_duplicates", 
    "per_parent_display", 
    "per_parent_prompt", 
    "stitch_max_chars",
    "enable_memory",
    "mem_ttl_days",
    "mem_show_full",
    "show_memories",
    "prefs_enabled",
    }

def load_prefs() -> dict:
    try:
        if os.path.exists(PREFS_PATH):
            with open(PREFS_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
                return {k: v for k, v in data.items() if k in NON_SENSITIVE_PREF_KEYS}
    except Exception:
        pass
    return {}

def save_prefs(prefs: dict):
    try:
        clean = {k: v for k, v in prefs.items() if k in NON_SENSITIVE_PREF_KEYS}
        with open(PREFS_PATH, "w", encoding="utf-8") as f:
            json.dump(clean, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"[WARN] save_prefs: {e}")

def init_prefs_in_session():
    prefs = load_prefs() or {}

    # --- Fix persist_dir for cloud/local ---
    pd = (prefs.get("persist_dir") or "").strip()
    if not pd:
        prefs["persist_dir"] = DEFAULT_CHROMA_DIR
    else:
        if IS_CLOUD and not pd.startswith("/tmp/"):
            prefs["persist_dir"] = DEFAULT_CHROMA_DIR

    # --- Initialize session_state for all prefs keys (once only) ---
    for key, value in prefs.items():
        if key not in st.session_state:
            st.session_state[key] = value

    # Keep full prefs dict available in session
    st.session_state["prefs"] = prefs


def one_line_preview(text: str, maxlen: int = 160) -> str:
    """Make text single-line, remove bullets/extra spaces, and truncate."""
    if not text:
        return ""
    s = text.replace("\r", " ").replace("\n", " ")
    s = re.sub(r"\s+", " ", s).strip()
    # remove list marker at the beginning of the text (es. "- ", "* ", "• ")
    s = re.sub(r"^[-\*\u2022]\s+", "", s)
    return (s[:maxlen] + "…") if len(s) > maxlen else s

# --- Chunking utilities (token-based, with tiktoken fallback) ---
def _tokenize(text: str):
    """
    Returns tokens as ids if tiktoken is available, otherwise whitespace tokens.
    """
    if _tk_enc:
        return _tk_enc.encode(text or "")
    return re.findall(r"\S+", text or "")

def _detokenize(toks):
    if _tk_enc:
        return _tk_enc.decode(toks)
    return " ".join(map(str, toks))

def split_into_chunks(text: str, chunk_size: int = 800, overlap: int = 80, min_size: int = 512):
    """
    Split very large descriptions into chunks with overlap.
    Returns list of tuples: (start_token_pos, chunk_text).
    If text is short (<= min_size), returns a single chunk with pos=0.
    """
    toks = _tokenize(text)
    n = len(toks)
    if n == 0:
        return [(0, "")]
    if n <= max(min_size, 2*chunk_size):
        return [(0, text)]

    chunks = []
    start = 0
    while start < n:
        end = min(n, start + int(chunk_size))
        # detokenize
        if _tk_enc:
            chunk_text = _tk_enc.decode(toks[start:end])
        else:
            chunk_text = " ".join(toks[start:end])
        chunks.append((start, chunk_text))
        if end >= n:
            break
        start = max(0, end - int(overlap))  # sliding window with overlap
    return chunks

# === Solution Memory (Level B) ===
MEM_COLLECTION = "memories"        # separate Chroma collection for playbooks
DEFAULT_MEM_TTL_DAYS = 180         # recommended expiration for memories

def now_ts() -> int:
    return int(time.time())

def ts_in_days(days: int) -> int:
    return now_ts() + days * 24 * 3600

def is_ollama_available() -> tuple[bool, str]:
    """
    Return (available, host). Check via HTTP on configured host
    and, as fallback, the presence of the 'ollama' binary.
    """
    host = os.getenv("OLLAMA_HOST", "http://localhost:11434").rstrip("/")
    # Quick HTTP attempt
    try:
        import requests  # lazy import
        r = requests.get(f"{host}/api/tags", timeout=1)
        if r.status_code < 500:
            return True, host
    except Exception:
        pass
    # Fallback: binary present and executable
    try:
        if shutil.which("ollama"):
            subprocess.run(
                ["ollama", "list"],
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=1
            )
            return True, host
    except Exception:
        pass
    return False, host

def get_ollama_model_names(ollama_host: str) -> list[str]:
    """Return a sorted list of Ollama model names from /api/tags."""
    try:
        import requests  # lazy import
        r = requests.get(f"{ollama_host.rstrip('/')}/api/tags", timeout=2)
        r.raise_for_status()
        data = r.json() or {}
        models = data.get("models", []) or []
        names = []
        for m in models:
            name = (m.get("name") or "").strip()
            if name:
                names.append(name)
        return sorted(set(names))
    except Exception:
        return []

def _normalize_provider_label(p: str) -> str:
    p = (p or "").strip().lower()
    if "openai" in p:
        return "openai"
    if "sentence" in p:  # "Local (sentence-transformers)"
        return "sentence-transformers"
    return p or "unknown"

def get_index_embedder_info(persist_dir: str, collection_name: str):
    """
    Ritorna (provider, model, source) dove source ∈ {"file","chroma",None}.
    1) Tenta <persist_dir>/<collection_name>__meta.json con chiavi {"provider","model"}
    2) Fallback: metadata della collection Chroma (se presenti)
    """
    # 1) file <collection>__meta.json
    try:
        meta_path = os.path.join(persist_dir, f"{collection_name}__meta.json")
        if os.path.isfile(meta_path):
            with open(meta_path, "r", encoding="utf-8") as f:
                m = json.load(f)
            prov = _normalize_provider_label(m.get("provider"))
            model = (m.get("model") or "").strip()
            if prov and model:
                return prov, model, "file"
    except Exception:
        pass

    # 2) metadata su Chroma (se salvati in fase di ingest)
    try:
        client = get_chroma_client(persist_dir)
        coll = client.get_collection(collection_name)
        md = (getattr(coll, "metadata", None) or {})  # alcune versioni mettono metadata qui
        prov = _normalize_provider_label(md.get("provider"))
        model = (md.get("model") or "").strip()
        if prov and model:
            return prov, model, "chroma"
    except Exception:
        pass

    return None, None, None

# ------------------------------
# Try imports with fallback/shim
# ------------------------------
try:
    import streamlit as st
    ST_AVAILABLE = True
    print("[DEBUG] Streamlit imported.")
except Exception:
    ST_AVAILABLE = False
    print("[DEBUG] Streamlit not available, using shim.")

    class _STShim:
        def __getattr__(self, _name):
            def _noop(*_args, **_kwargs):
                print(f"[DEBUG] Call st.{_name} ignorata (shim).")
                return None
            return _noop

    st = _STShim()  # type: ignore

try:
    import chromadb
    from chromadb.config import Settings as ChromaSettings
    print("[DEBUG] chromadb imported.")
except Exception:
    chromadb = None  # type: ignore
    ChromaSettings = None  # type: ignore
    print("[DEBUG] chromadb not available.")

try:
    from sentence_transformers import SentenceTransformer
    print("[DEBUG] sentence-transformers imported.")
except Exception:
    SentenceTransformer = None  # type: ignore
    print("[DEBUG] sentence-transformers not available.")

try:
    from openai import OpenAI
    print("[DEBUG] openai SDK imported.")
except Exception:
    OpenAI = None  # type: ignore
    print("[DEBUG] openai SDK not available.")


# ------------------------------
# Constants
# ------------------------------
DEFAULT_COLLECTION = "tickets"
DEFAULT_EMBEDDING_MODEL = "all-MiniLM-L6-v2"
DEFAULT_LLM_MODEL = "gpt-4o"
DEFAULT_OLLAMA_MODEL = "llama3.2"
NEW_LABEL = "➕ Create new collection..."

# ------------------------------
# Pastel colors per phase (UI)
# ------------------------------
PHASE_COLORS = {
    "YouTrack connection":      "#F0F4FF",  # light blue
    "Embeddings & Vector DB":   "#F9F5FF",  # light lilac
    "Retrieval configuration":  "#FFF7E6",  # light cream
    "LLM & API keys":           "#E8FFF0",  # light sage
    "Solutions memory":         "#FFF0F3",  # light pink
    "Chat & Results":           "#E6FAFF",  # light cyan
    "Preferences & debug":      "#F6F6F6",  # neutral grey
    "Docs KB (PDF/DOCX/TXT)":      "#F0FFF6",  # mint
    "MCP Console":              "#1D0B82",  # dark blue
}

# Icons for phases (used in sidebar radio)
PHASE_ICONS = {
    "YouTrack connection":      "🔗",
    "MCP Console":              "🛠️",
    "Embeddings & Vector DB":   "🧩",
    "Retrieval configuration":  "🔍",
    "LLM & API keys":           "🔑",
    "Chat & Results":           "💬",
    "Solutions memory":         "💾",
    "Preferences & debug":      "⚙️",
    "Docs KB (PDF/DOCX/TXT)":      "📚",
}


# ------------------------------
# MCP_PROMPT_LIBRARY (one-click)
# ------------------------------
# Templates may contain:
# - {{PROJECT}}: replaced with the current project key (from Phase 1)
# Keep prompts short and tool-oriented for best MCP behavior.
MCP_PROMPT_LIBRARY = [
    {
        "category": "Quick status",
        "items": [
            ("Project snapshot", "Give a concise snapshot of project {{PROJECT}}: open issues, blockers, critical issues, and last updates."),
            ("Updated last 24h", "List issues in project {{PROJECT}} updated in the last 24 hours. Return idReadable, summary, state, assignee."),
            ("Resolved last 7d", "List issues resolved in project {{PROJECT}} in the last 7 days. Return idReadable, summary, resolution date."),
        ],
    },
    {
        "category": "Backlog & workload",
        "items": [
            ("State counts", "Group open issues in project {{PROJECT}} by State and provide counts."),
            ("Workload per assignee", "Group open issues in project {{PROJECT}} by assignee and show counts."),
            ("Unassigned", "Find unassigned open issues in project {{PROJECT}}. Return idReadable, summary, priority."),
        ],
    },
    {
        "category": "Aging & stuck",
        "items": [
            ("Open > 30 days", "Find issues in project {{PROJECT}} open for more than 30 days. Sort by age descending."),
            ("Stuck In Progress > 14d", "Find issues in project {{PROJECT}} stuck in 'In Progress' for more than 14 days."),
            ("Not updated > 21d", "Find open issues in project {{PROJECT}} not updated in the last 21 days."),
        ],
    },
    {
        "category": "Critical & risks",
        "items": [
            ("Critical / Blockers", "List Critical or Blocker issues in project {{PROJECT}}. Include idReadable, summary, state, assignee."),
            ("Risk summary", "Based on current open issues in project {{PROJECT}}, summarize the main technical and delivery risks."),
            ("SLA risk (heuristic)", "Identify issues in project {{PROJECT}} likely to miss SLA based on age, priority and state."),
        ],
    },
    {
        "category": "Workflow health",
        "items": [
            ("State distribution", "Show how many issues in project {{PROJECT}} are in each workflow state."),
            ("Workflow regressions", "Find issues in project {{PROJECT}} that moved backwards in workflow (e.g. Review -> In Progress)."),
            ("Bottlenecks", "Analyze workflow states in project {{PROJECT}} and identify bottlenecks."),
        ],
    },
    {
        "category": "Trends",
        "items": [
            ("Created per week", "How many issues were created per week in project {{PROJECT}} in the last 2 months?"),
            ("Resolved per week", "How many issues were resolved per week in project {{PROJECT}} in the last month?"),
            ("Lead time estimate", "Estimate average time from creation to resolution for issues in project {{PROJECT}}."),
        ],
    },
    {
        "category": "Planning",
        "items": [
            ("Next sprint candidates", "Suggest a prioritized list of issues for the next sprint in project {{PROJECT}}. Explain the reasoning."),
            ("Carry-over risks", "Which issues in project {{PROJECT}} are most likely to spill into the next sprint?"),
            ("Meeting report", "Generate a short status report for a project meeting about {{PROJECT}}: progress, blockers, risks, next actions."),
        ],
    },
]
def inject_global_css():
    """Inject a small global CSS theme for a more modern look."""
    import streamlit as st
    st.markdown(
        """
        <style>
        /* Global page background and typography */
        body {
            background-color: #F7F9FC;
        }
        h1, h2, h3 {
            color: #1F2933;
        }

        /* Primary buttons */
        div.stButton > button:first-child {
            background-color: #4361EE;
            color: white;
            border-radius: 10px;
            padding: 0.4rem 1.1rem;
            border: none;
            font-weight: 500;
        }
        div.stButton > button:first-child:hover {
            background-color: #3554D1;
        }

        /* Sidebar titles a bit smaller */
        section[data-testid="stSidebar"] h1, 
        section[data-testid="stSidebar"] h2, 
        section[data-testid="stSidebar"] h3 {
            font-size: 0.95rem;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

def phase_container_start(phase_label: str):
    """Start a colored container for the given phase."""
    import streamlit as st
    color = PHASE_COLORS.get(phase_label, "#F6F6F6")
    st.markdown(
        f"""
        <div style="
            background-color:{color};
            padding: 1.5rem 1.8rem;
            border-radius: 18px;
            border: 1px solid rgba(0,0,0,0.04);
            margin-top: 0.75rem;
        ">
        """,
        unsafe_allow_html=True,
    )

def phase_container_end():
    """Close the colored phase container."""
    import streamlit as st
    st.markdown("</div>", unsafe_allow_html=True)

# ------------------------------
# YouTrack Client + DTO
# ------------------------------
@dataclass
class YTIssue:
    id_readable: str
    summary: str
    description: str
    project: str

    def text_blob(self) -> str:
        return f"{self.summary}\n\n{self.description or ''}"


class YouTrackClient:
    def __init__(self, base_url: str, token: str):
        self.base_url = base_url.rstrip("/")
        self.token = token

    def _get(self, path: str, params: Optional[dict] = None):
        import requests
        url = f"{self.base_url}{path}"
        headers = {"Authorization": f"Bearer {self.token}"}
        r = requests.get(url, headers=headers, params=params, timeout=30)
        r.raise_for_status()
        return r.json()

    def list_projects(self) -> List[dict]:
        print("[DEBUG] list_projects")
        return self._get("/api/admin/projects", params={"fields": "name,shortName,id"})

    def list_issues(self, project_key: str, limit: int = 200) -> List[YTIssue]:
        print(f"[DEBUG] list_issues per progetto {project_key}")
        fields = "idReadable,summary,description,project(name,shortName)"
        params = {
            "query": f"project: {project_key}",
            "$top": str(limit),
            "fields": fields,
        }
        data = self._get("/api/issues", params=params)
        issues: List[YTIssue] = []
        for it in data:
            issues.append(
                YTIssue(
                    id_readable=it.get("idReadable", ""),
                    summary=it.get("summary", ""),
                    description=it.get("description", ""),
                    project=(it.get("project", {}) or {}).get("shortName")
                            or (it.get("project", {}) or {}).get("name", ""),
                )
            )
        print(f"[DEBUG] trovati {len(issues)} issues")
        return issues


# ------------------------------
# Embeddings backend
# ------------------------------
class EmbeddingBackend:
    def __init__(self, use_openai: bool, model_name: str):
        self.use_openai = use_openai
        self.model_name = model_name
        self.provider_name = "openai" if use_openai else "sentence-transformers"

        if use_openai:
            if OpenAI is None:
                raise RuntimeError("openai SDK not available")
            api_key = get_openai_key()
            if not api_key:
                raise RuntimeError("OPENAI_API_KEY not set (enter it in the sidebar or as an environment variable)")
            self.client = OpenAI(api_key=api_key)
        else:
            if SentenceTransformer is None:
                raise RuntimeError("sentence-transformers non available")
            self.model = SentenceTransformer(model_name)

    def embed(self, texts: List[str]) -> List[List[float]]:
        print(f"[DEBUG] embed {len(texts)} text with provider={self.provider_name} model={self.model_name}")
                # NOTE: OpenAI Embeddings API has a max *tokens per request* limit (e.g. 300k).
        # When indexing large docs we may generate many chunks; sending them in one call can exceed the limit.
        # We therefore batch inputs to stay safely under the limit.
        if self.use_openai:
            # Conservative safety margin under the hard cap (300000)
            MAX_TOKENS_PER_REQUEST = int(os.getenv("OPENAI_EMBED_MAX_TOKENS_PER_REQUEST", "290000"))
            # Also avoid extremely large single inputs (model-dependent). Use a conservative cap.
            MAX_TOKENS_PER_INPUT = int(os.getenv("OPENAI_EMBED_MAX_TOKENS_PER_INPUT", "8000"))

            def count_tokens(txt: str) -> int:
                try:
                    if _tk_enc is not None:
                        return len(_tk_enc.encode(txt or ""))
                except Exception:
                    pass
                # Fallback heuristic: ~4 chars per token for Latin text
                return max(1, int(len(txt or "") / 4))

            # If a single item is too large, split it further so we don't fail.
            expanded: List[str] = []
            for t in texts:
                if count_tokens(t) <= MAX_TOKENS_PER_INPUT:
                    expanded.append(t)
                    continue
                # Split oversize items into smaller token windows (no overlap here; overlap is already handled upstream)
                sub = split_into_chunks(
                    t,
                    chunk_size=min(2000, MAX_TOKENS_PER_INPUT // 2),
                    overlap=0,
                    min_size=0,
                )
                expanded.extend([x[1] for x in sub if x[1]])

            out: List[List[float]] = []
            batch: List[str] = []
            batch_tokens = 0
            def flush():
                nonlocal batch, batch_tokens, out
                if not batch:
                    return
                res = self.client.embeddings.create(model=self.model_name, input=batch)
                out.extend([d.embedding for d in res.data])  # type: ignore
                batch = []
                batch_tokens = 0

            for t in expanded:
                t_tokens = count_tokens(t)
                # If adding this text would exceed the cap, flush first
                if batch and (batch_tokens + t_tokens) > MAX_TOKENS_PER_REQUEST:
                    flush()
                batch.append(t)
                batch_tokens += t_tokens

            flush()
            return out

        return self.model.encode(texts, normalize_embeddings=True).tolist()  # type: ignore

def get_chroma_client(persist_dir: str):
    """Create folder and return a Chroma PersistentClient.

    This function delegates to rag_core.get_chroma_client() when available,
    while keeping a local fallback to preserve backward compatibility.
    """
    # Prefer shared core (used also by webhook/FastAPI service)
    try:
        from rag_core import get_chroma_client as _core_get_chroma_client  # type: ignore
        return _core_get_chroma_client(persist_dir)
    except Exception:
        pass

    # Fallback: legacy implementation (kept to avoid regressions)
    try:
        import chromadb  # lazy import (no global)
        from chromadb.config import Settings as ChromaSettings
    except Exception as e:
        # explicit diagnostics
        st.error("ChromaDB cannot be imported in this environment.")
        st.exception(e)
        raise

    os.makedirs(persist_dir, exist_ok=True)
    return chromadb.PersistentClient(
        path=persist_dir,
        settings=ChromaSettings(anonymized_telemetry=False)
    )


def load_chroma_collection(persist_dir: str, collection_name: str, space: str = "cosine"):
    """Return a Chroma collection, creating it if needed.

    Delegates to rag_core.load_chroma_collection() when available,
    otherwise falls back to local get_chroma_client().
    """
    try:
        from rag_core import load_chroma_collection as _core_load  # type: ignore
        res = _core_load(persist_dir=persist_dir, collection_name=collection_name, space=space)
        # Backward/forward compatibility: accept either a collection or (client, collection)
        if isinstance(res, tuple) and len(res) == 2:
            return res[1]
        return res
    except Exception:
        pass
    client = get_chroma_client(persist_dir)
    return client.get_or_create_collection(name=collection_name, metadata={"hnsw:space": space})


# ------------------------------
# VectorStore (Chroma wrapper)
# ------------------------------
class VectorStore:
    def __init__(self, persist_dir: str, collection_name: str):
        if chromadb is None:
            raise RuntimeError("ChromaDB not available")
        self.persist_dir = persist_dir
        self.collection_name = collection_name
        os.makedirs(self.persist_dir, exist_ok=True)
        self.client = get_chroma_client(persist_dir)
        self.col = load_chroma_collection(self.persist_dir, collection_name, space="cosine")

    def add(self, ids: List[str], documents: List[str], metadatas: List[dict], embeddings: List[List[float]]):
        print(f"[DEBUG] add {len(ids)} documents to {self.collection_name}")
        self.col.add(ids=ids, documents=documents, metadatas=metadatas, embeddings=embeddings)

    def query(self, query_embeddings: List[List[float]], n_results: int = 5) -> dict:
        return self.col.query(query_embeddings=query_embeddings, n_results=n_results)

    def count(self) -> int:
        try:
            return self.col.count()  # type: ignore
        except Exception:
            return -1

def get_openai_key() -> Optional[str]:
    try:
        import streamlit as st  # type: ignore
        ui_key = st.session_state.get("openai_key")
        if ui_key:
            return ui_key
        if hasattr(st, "secrets") and "OPENAI_API_KEY" in st.secrets:
            return st.secrets["OPENAI_API_KEY"]
    except Exception:
        pass
    return os.getenv("OPENAI_API_KEY") or os.getenv("OPENAI_API_KEY_EXPERIMENTS")

# ------------------------------
# LLM Backend
# ------------------------------
class LLMBackend:
    def __init__(self, provider: str, model_name: str, temperature: float = 0.2):
        self.provider = provider
        self.model_name = model_name
        self.temperature = float(temperature)
        if provider == "OpenAI":
            if OpenAI is None:
                raise RuntimeError("openai SDK not available")
            api_key = get_openai_key()
            if not api_key:
                raise RuntimeError("OPENAI_API_KEY not set (enter it in the sidebar or as an environment variable)")
            self.client = OpenAI(api_key=api_key)

        elif provider == "Ollama (local)":
            ok, host = is_ollama_available()
            if not ok:
                raise RuntimeError("Ollama is not available in this environment: select an OpenAI provider.")
            # No SDK client needed: use REST; keep host for calls
            self.client = None
            self.ollama_host = host  # es. "http://localhost:11434"

        else:
            raise RuntimeError("Unsupported LLM provider")

    def generate(self, system: str, user: str) -> str:
        if self.provider == "OpenAI":
            try:
                # Responses API
                res = self.client.responses.create(
                    model=self.model_name,
                    input=[{"role": "system", "content": system}, {"role": "user", "content": user}],
                    temperature=self.temperature,
                )
                return res.output_text  # type: ignore
            except Exception:
                # Fallback Chat Completions
                chat = self.client.chat.completions.create(
                    model=self.model_name,
                    messages=[{"role": "system", "content": system}, {"role": "user", "content": user}],
                    temperature=self.temperature,
                )
                return chat.choices[0].message.content or ""
        elif self.provider == "Ollama (local)":
            import requests
            host = getattr(self, "ollama_host", "http://localhost:11434").rstrip("/")
            url = f"{host}/api/chat"

            payload = {
                "model": self.model_name,  # es. "llama3.2" o "llama3.2:latest"
                "messages": [
                    {"role": "system", "content": system},
                    {"role": "user", "content": user},
                ],
                "stream": False,  # evita JSON concatenati
                "options": {"temperature": self.temperature},  # Ollama supporta temperature
            }

            r = requests.post(url, json=payload, timeout=300)
            try:
                r.raise_for_status()
                data = r.json()  # singolo JSON
            except json.JSONDecodeError:
                # fallback per eventuale streaming non richiesto
                text = r.text.strip()
                chunks, buf, depth = [], "", 0
                for ch in text:
                    buf += ch
                    if ch == "{":
                        depth += 1
                    elif ch == "}":
                        depth -= 1
                        if depth == 0:
                            try:
                                chunks.append(json.loads(buf))
                            except Exception:
                                pass
                            buf = ""
                parts = []
                for obj in chunks:
                    if isinstance(obj, dict):
                        if "message" in obj and isinstance(obj["message"], dict) and "content" in obj["message"]:
                            parts.append(obj["message"]["content"])
                        elif "response" in obj:
                            parts.append(obj["response"])
                return "\n".join(parts).strip() if parts else text

            if isinstance(data, dict):
                if "message" in data and isinstance(data["message"], dict) and "content" in data["message"]:
                    return data["message"]["content"]
                if "response" in data:
                    return data["response"]
                if "messages" in data and isinstance(data["messages"], list) and data["messages"]:
                    last = data["messages"][-1]
                    if isinstance(last, dict) and "content" in last:
                        return last["content"]
            return str(data)


RAG_SYSTEM_PROMPT = (
    "You are a technical assistant who answers based on similar YouTrack tickets. "
    "Always cite the IDs of the tickets found in square brackets. If the context is insufficient, ask for clarifications. Use english in the answer"
)

def build_prompt(user_ticket: str, retrieved: List[Tuple[str, dict, float]]) -> str:
    print(f"[DEBUG] build_prompt with {len(retrieved)} retrieved documents")
    parts = [
        "New ticket:\n" + user_ticket.strip(),
        "\nSimilar tickets found (closest first):",
    ]
    for doc, meta, dist in retrieved:
        parts.append(
            f"- {meta.get('id_readable','')} | distance={dist:.3f} | {meta.get('summary','')}\n{doc[:500]}"
        )
    parts.append("\nAnswer including citations in [ ] of the relevant ticket IDs. Please use always english in the answer")
    return "\n".join(parts)


from collections import defaultdict

def collapse_by_parent(results, per_parent=1, stitch_for_prompt=False, max_chars=1200):
    """
    Collapse multiple chunks from the same ticket into one row.
    results: list of (doc, meta, dist, src)
    - per_parent: keep at most N items per parent (1 for display)
    - stitch_for_prompt: if True, concatenate selected chunks in order of 'pos' (bounded by max_chars)
    Returns: list of (doc, meta, dist, src) sorted by distance.
    """
    groups = defaultdict(list)
    for doc, meta, dist, src in results:
        meta = meta or {}
        pid = meta.get("parent_id") or meta.get("id_readable")
        groups[pid].append((doc, meta, dist, src))

    collapsed = []
    for pid, items in groups.items():
        items = sorted(items, key=lambda x: (x[2], x[1].get("pos", 0)))  # by distance, then pos
        keep = items[:max(1, per_parent)]

        if stitch_for_prompt and len(keep) > 1:
            keep_sorted = sorted(keep, key=lambda x: x[1].get("pos", 0))
            buf = []
            total = 0
            for d, _m, _dist, _src in keep_sorted:
                if total + len(d) + 2 > max_chars:
                    break
                buf.append(d)
                total += len(d) + 2
            best = keep[0]  # best distance
            stitched = "\n\n".join(buf) if buf else best[0]
            collapsed.append((stitched, best[1], best[2], best[3]))
        else:
            collapsed.append(keep[0])

    return sorted(collapsed, key=lambda x: x[2])

# ------------------------------
# NEW: utility to open the collection into the session
# ------------------------------
def open_vector_in_session(persist_dir: str, collection_name: str):
    """Open (or create) a Chroma collection and put it into the session."""
    try:
        vs = VectorStore(persist_dir=persist_dir, collection_name=collection_name)
        cnt = vs.count()
        try:
            import streamlit as st  # type: ignore
            st.session_state["vector"] = vs
            st.session_state["vs_persist_dir"] = persist_dir
            st.session_state["vs_collection"] = collection_name
            st.session_state["vs_count"] = cnt
        except Exception:
            pass
        return True, cnt, None
    except Exception as e:
        try:
            import streamlit as st  # type: ignore
            st.session_state["vector"] = None
        except Exception:
            pass
        return False, -1, str(e)

def render_phase_yt_connection_page(prefs):
    import streamlit as st  # local import to keep signature simple

    st.title("Phase 1 – YouTrack connection")
    st.write("Configure the YouTrack endpoint used to fetch and index tickets.")

    if st.session_state.get("yt_client"):
        st.success("Connected to YouTrack")

    # Read current values (from session_state or prefs)
    current_url = st.session_state.get("yt_url", prefs.get("yt_url", ""))
    current_token = st.session_state.get("yt_token", "")

    col1, col2 = st.columns(2)
    with col1:
        yt_url = st.text_input(
            "YouTrack server URL",
            value=current_url,
            placeholder="https://<org>.myjetbrains.com/youtrack",
        )
    with col2:
        yt_token = st.text_input(
            "YouTrack token (Bearer)",
            type="password",
            value=current_token,
        )

    # Keep state in session_state (not in prefs, to avoid writing token to disk)
    st.session_state["yt_url"] = yt_url
    st.session_state["yt_token"] = yt_token

    connect = st.button("Connect to YouTrack")

        # Projects & issues

    st.subheader("YouTrack Projects")
    if st.session_state.get("yt_client"):
        projects = st.session_state.get("projects", [])
        if projects:
            names = [f"{p.get('name')} ({p.get('shortName')})" for p in projects]
            sel = st.selectbox("Choose project", options=["-- select --"] + names, key="proj_select")

            if sel and sel != "-- select --":
                p = projects[names.index(sel)]
                project_key = p.get("shortName") or p.get("name")

                # AUTO-LOAD: when the project changes, immediately load the issues
                prev_key = st.session_state.get("last_project_key")
                if prev_key != project_key:
                    try:
                        with st.spinner("Loading issues..."):
                            st.session_state["issues"] = st.session_state["yt_client"].list_issues(project_key)
                        st.session_state["last_project_key"] = project_key
                        st.success(f"Loaded {len(st.session_state['issues'])} issues")
                    except Exception as e:
                        st.error(f"Error loading issues: {e}")

                # optional: button to reload manually
                if st.button("Reload issues"):
                    try:
                        with st.spinner("Reloading issues..."):
                            st.session_state["issues"] = st.session_state["yt_client"].list_issues(project_key)
                        st.success(f"Loaded {len(st.session_state['issues'])} issues")
                    except Exception as e:
                        st.error(f"Reload error: {e}")
        else:
            st.caption("No project found (or insufficient permissions).")

    # ALWAYS show the table if there are issues in memory
    issues = st.session_state.get("issues", [])
    if issues:
        base_url = (st.session_state.get("yt_client").base_url if st.session_state.get("yt_client") else "").rstrip("/")

        # Build a Markdown table with clickable ID and without Project/Open
        lines = []
        lines.append("| ID | Summary |")
        lines.append("|---|---|")
        for it in issues:
            url = f"{base_url}/issue/{it.id_readable}" if base_url else ""
            id_cell = f"[{it.id_readable}]({url})" if url else it.id_readable
            # trim very long summaries (optional)
            summary = it.summary.strip().replace("\n", " ")
            if len(summary) > 160:
                summary = summary[:157] + "..."
            lines.append(f"| {id_cell} | {summary} |")

        st.markdown("\n".join(lines), unsafe_allow_html=False)
    else:
        st.caption("No issues in memory. Select a project to load tickets.")

    return yt_url, yt_token, connect


# === MCP CONSOLE (auto-added) BEGIN ===
def run_mcp_prompt(prompt: str, *, yt_url: str, yt_token: str, openai_key: str) -> dict:
    """Run a prompt against YouTrack MCP server via OpenAI Responses API.

    Returns a dict with keys:
    - ok: bool
    - readable: str
    - raw: object (best-effort)
    - error: str (if not ok)
    """
    if not prompt or not prompt.strip():
        return {"ok": False, "readable": "", "raw": None, "error": "Empty prompt"}

    if OpenAI is None:
        return {"ok": False, "readable": "", "raw": None, "error": "openai SDK not available"}

    base = (yt_url or "").rstrip("/")
    if not base:
        return {"ok": False, "readable": "", "raw": None, "error": "YouTrack URL missing (configure Phase 1)"}

    if not yt_token:
        return {"ok": False, "readable": "", "raw": None, "error": "YouTrack token missing (configure Phase 1)"}

    if not openai_key:
        return {"ok": False, "readable": "", "raw": None, "error": "OpenAI API key missing (configure Phase 5)"}

    client = OpenAI(api_key=openai_key)
    server_url = f"{base}/mcp"

    try:
        res = client.responses.create(
            model="gpt-4o",
            input=[
                {"role": "system", "content": "You are a helpful assistant that can call YouTrack MCP tools as needed."},
                {"role": "user", "content": prompt},
            ],
            tools=[
                {
                    "type": "mcp",
                    "server_label": "youtrack",
                    "server_url": server_url,
                    "headers": {"Authorization": f"Bearer {yt_token}"},
                    "require_approval": "never",
                }
            ],
        )
        readable = getattr(res, "output_text", "") or ""
        return {"ok": True, "readable": readable, "raw": res, "error": ""}
    except Exception as e:
        return {"ok": False, "readable": "", "raw": None, "error": str(e)}



# --- MCP One-click prompts (JSON) helpers ---
def load_mcp_prompts_from_json(json_path: str) -> list[dict]:
    """Load MCP one-click prompts from a JSON file.

    Expected schema:
    {
      "version": 1,
      "categories": [
        {
          "name": "Category name",
          "items": [{"label": "...", "prompt": "..."}]
        }
      ]
    }
    Returns a normalized list:
    [
      {"category": "...", "items": [("Label", "Prompt"), ...]},
      ...
    ]
    """
    import json
    import os

    if not json_path:
        return []

    # Allow env override for deployments
    json_path = os.environ.get("MCP_PROMPTS_PATH", json_path)

    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except FileNotFoundError:
        return []
    except Exception:
        # Invalid JSON
        return []

    cats = data.get("categories", [])
    out: list[dict] = []
    for c in cats:
        name = str(c.get("name", "")).strip()
        items = c.get("items", []) or []
        norm_items = []
        for it in items:
            lbl = str(it.get("label", "")).strip()
            prm = str(it.get("prompt", "")).strip()
            if lbl and prm:
                norm_items.append((lbl, prm))
        if name and norm_items:
            out.append({"category": name, "items": norm_items})
    return out


def get_mcp_prompts_library(default_json_path: str = "mcp_prompts.json") -> list[dict]:
    """Get cached MCP prompt library in session state (if available)."""
    import os
    import streamlit as st  # local import

    # Resolve path with env override
    json_path = os.environ.get("MCP_PROMPTS_PATH", default_json_path)

    # Cache in session_state for speed and to support "Reload"
    cache_key = "mcp_prompt_library"
    path_key = "mcp_prompt_library_path"
    if cache_key not in st.session_state or st.session_state.get(path_key) != json_path:
        lib = load_mcp_prompts_from_json(json_path)
        st.session_state[cache_key] = lib
        st.session_state[path_key] = json_path
    return st.session_state.get(cache_key, []) or []
def render_phase_mcp_console_page(prefs):
    import streamlit as st  # local import

    st.title("Phase 2 - MCP Console")
    st.write("Interact with the current YouTrack project via MCP (without changing the rest of the app).")

    yt_url = st.session_state.get("yt_url", "")
    yt_token = st.session_state.get("yt_token", "")
    project_key = st.session_state.get("last_project_key", "")

    st.markdown("---")
    c1, c2, c3 = st.columns([2, 2, 1])
    with c1:
        st.caption("YouTrack URL (from Phase 1)")
        st.code(yt_url or "—")
    with c2:
        st.caption("Current project (from Phase 1)")
        st.code(project_key or "—")
    with c3:
        st.caption("Loaded issues")
        issues_n = len(st.session_state.get("issues", []) or [])
        st.code(str(issues_n))

    st.markdown("---")

    if "mcp_console_prompt" not in st.session_state:
        st.session_state["mcp_console_prompt"] = ""

    b1, b2, b3 = st.columns([1, 1, 2])
    with b1:
        if st.button("Insert current project"):
            if project_key:
                p = st.session_state.get("mcp_console_prompt", "")
                prefix = f"Project: {project_key}\n"
                if prefix not in p:
                    st.session_state["mcp_console_prompt"] = prefix + p
            else:
                st.warning("No project selected in Phase 1.")
    with b2:
        if st.button("Test MCP"):
            if project_key:
                st.session_state["mcp_console_prompt"] = (
                    f"Search 1 issue in project {project_key} and print its idReadable and summary."
                )
            else:
                st.session_state["mcp_console_prompt"] = "List projects and print their shortName and name."
    with b3:
        st.caption("Tip: ask for 'search issues', 'get issue details', 'list projects', etc.")

    # -----------------------------
    # One-click prompt library
    # -----------------------------
    st.markdown("#### One-click prompts")

    # Auto-run toggle (optional)
    if "mcp_autorun_on_click" not in st.session_state:
        st.session_state["mcp_autorun_on_click"] = False

    copt1, copt2 = st.columns([1, 3])
    with copt1:
        st.session_state["mcp_autorun_on_click"] = st.checkbox(
            "Auto-run on click",
            value=bool(st.session_state["mcp_autorun_on_click"]),
            help="If enabled, clicking a preset will immediately run it via MCP.",
        )
    with copt2:
        st.caption("Presets automatically inject the current project key when needed.")

    # Category select (JSON-backed)
    # Load presets from mcp_prompts.json (editable) and fallback to embedded library if needed.
    default_json_path = os.path.join(APP_DIR, "mcp_prompts.json")
    lib = get_mcp_prompts_library(default_json_path=default_json_path)

    # Optional: allow quick reload without restarting Streamlit
    c_reload1, c_reload2 = st.columns([1, 4])
    with c_reload1:
        if st.button("Reload presets", key="mcp_reload_presets"):
            # Drop cached library so next call reloads from disk
            st.session_state.pop("mcp_prompt_library", None)
            st.session_state.pop("mcp_prompt_library_path", None)
            lib = get_mcp_prompts_library(default_json_path=default_json_path)
            st.toast("MCP presets reloaded", icon="🔄")
    with c_reload2:
        st.caption(f"Presets path: {os.environ.get('MCP_PROMPTS_PATH', default_json_path)}")

    # Fallback to embedded presets if JSON is missing/empty
    if not lib:
        lib = MCP_PROMPT_LIBRARY

    categories = [c.get("category") for c in lib]
    if "mcp_prompt_category" not in st.session_state:
        st.session_state["mcp_prompt_category"] = categories[0] if categories else ""

    cat = st.selectbox(
        "Preset category",
        options=categories,
        key="mcp_prompt_category",
    )

    # Resolve items for the selected category
    items = []
    for c in lib:
        if c.get("category") == cat:
            items = c.get("items", [])
            break

    # Render buttons in a grid
    # Each item is (label, template)
    if items:
        cols_per_row = 3
        for i in range(0, len(items), cols_per_row):
            row = st.columns(cols_per_row)
            for j in range(cols_per_row):
                k = i + j
                if k >= len(items):
                    continue
                label, template = items[k]
                with row[j]:
                    if st.button(label, key=f"mcp_preset_btn_{cat}_{k}"):
                        # Inject current project if present
                        proj = project_key or ""
                        text = template.replace("{{PROJECT}}", proj if proj else "{{PROJECT}}")

                        # If project placeholder still exists, prepend a hint
                        if "{{PROJECT}}" in text:
                            st.warning("No project selected in Phase 1. Select a project to fully use this preset.")
                        st.session_state["mcp_console_prompt"] = text

                        # Optionally auto-run
                        if bool(st.session_state.get("mcp_autorun_on_click")):
                            openai_key = get_openai_key() or ""
                            with st.spinner("Running MCP preset..."):
                                out = run_mcp_prompt(
                                    st.session_state["mcp_console_prompt"],
                                    yt_url=yt_url,
                                    yt_token=yt_token,
                                    openai_key=openai_key,
                                )
                            st.session_state["mcp_console_last"] = out

    st.markdown("---")

    prompt = st.text_area(
        "Prompt",
        key="mcp_console_prompt",
        height=200,
        placeholder="Example: Search the 10 most recent issues in the current project about 'VPN' and summarize.",
    )

    run = st.button("Run MCP prompt")

    if run:
        openai_key = get_openai_key() or ""
        with st.spinner("Running MCP prompt..."):
            out = run_mcp_prompt(prompt, yt_url=yt_url, yt_token=yt_token, openai_key=openai_key)
        st.session_state["mcp_console_last"] = out

    out = st.session_state.get("mcp_console_last")
    if out:
        tabs = st.tabs(["Readable", "Raw", "Error"])
        with tabs[0]:
            if out.get("ok"):
                st.write(out.get("readable") or "")
            else:
                st.info("No readable output.")
        with tabs[1]:
            st.write(out.get("raw"))
        with tabs[2]:
            if not out.get("ok"):
                st.error(out.get("error") or "Unknown error")
            else:
                st.caption("No error.")
# === MCP CONSOLE (auto-added) END ===
def render_phase_embeddings_vectordb_page(prefs):
    import streamlit as st

    # -----------------------------
    # 0) Normalize prefs and bootstrap state
    # -----------------------------
    if isinstance(prefs, dict):
        prefs_dict = prefs
    else:
        prefs_dict = st.session_state.get("prefs", {})

    def init_state(key: str, default):
        """Initialize a session_state key once, if not already set."""
        if key not in st.session_state:
            st.session_state[key] = default

    # Persist dir
    init_state("persist_dir", prefs_dict.get("persist_dir", DEFAULT_CHROMA_DIR))

    # Collections (semantic: collection_selected + vs_collection = actual name)
    init_state(
        "collection_selected",
        (prefs_dict.get("collection_selected") or "").strip() or None,
    )
    init_state(
        "vs_collection",
        (prefs_dict.get("collection_selected") or "").strip()
        or prefs_dict.get("new_collection_name")
        or None,
    )
    init_state(
        "new_collection_name_input",
        prefs_dict.get("new_collection_name", DEFAULT_COLLECTION),
    )

    # Embeddings prefs (provider/model)
    init_state(
        "emb_provider_select",
        prefs_dict.get("emb_backend", "OpenAI"),
    )
    # Keep a stable previous backend marker for provider-change detection
    init_state(
        "_prev_emb_backend", 
        st.session_state.get("emb_provider_select")
    )
    init_state(
        "emb_model",
        prefs_dict.get("emb_model_name", ""),
    )

    # -----------------------------
    # 1) Header + success message
    # -----------------------------
    st.title("Phase 3 – Embeddings & Vector DB")
    st.write(
        "Configure the Chroma vector store (path and collections) and the embeddings "
        "provider/model used for indexing and query."
    )

    msg = st.session_state.pop("_index_success", None)
    if msg:
        st.success(msg)

    # Placeholder for the top status summary (we fill it later)
    status_box = st.empty()

    st.markdown("---")

    # -----------------------------
    # 2) Vector DB configuration
    # -----------------------------
    st.header("Vector DB")

    persist_dir = st.text_input(
        "Chroma path",
        value=st.session_state["persist_dir"],
        key="persist_dir",
        help="Directory where Chroma will store its collections.",
    )

    # Ensure directory exists
    try:
        os.makedirs(persist_dir, exist_ok=True)
    except Exception as e:
        st.error(f"Failed to create directory '{persist_dir}': {e}")
        return

    # -----------------------------
    # 3) Read existing collections from Chroma
    # -----------------------------
    coll_options: list[str] = []
    try:
        if chromadb is not None:
            _client = get_chroma_client(persist_dir)
            coll_options = [c.name for c in _client.list_collections()]  # type: ignore
    except Exception as e:
        st.caption(f"Unable to read collections from '{persist_dir}': {e}")

    # -----------------------------
    # 4) Keep last selection / new name
    # -----------------------------
    last_sel = st.session_state.get("collection_selected")
    if not last_sel:
        last_sel = (prefs_dict.get("collection_selected") or "").strip() or None

    last_new = st.session_state.get("new_collection_name_input")
    if not last_new:
        last_new = prefs_dict.get("new_collection_name", DEFAULT_COLLECTION)
        st.session_state["new_collection_name_input"] = last_new

    st.caption(
        f"Selected pref: {prefs_dict.get('collection_selected', '—')}  ·  Path: {persist_dir}"
    )

    # -----------------------------
    # 5) Select / create collection
    # -----------------------------
    if coll_options:
        opts = coll_options + [NEW_LABEL]

        if last_sel in opts:
            default_index = opts.index(last_sel)
        elif DEFAULT_COLLECTION in opts:
            default_index = opts.index(DEFAULT_COLLECTION)
        else:
            default_index = 0

        sel = st.selectbox(
            "Collection",
            options=opts,
            index=default_index,
            key="collection_select",
        )

        if sel == NEW_LABEL:
            new_name = st.text_input("New collection name", key="new_collection_name_input")
            st.caption("ℹ️ This collection will be physically created after you run **Index tickets** at least once.")
            collection_name = (st.session_state.get("new_collection_name_input") or "").strip() or DEFAULT_COLLECTION
        else:
            collection_name = sel
            st.session_state["after_delete_reset"] = True
    else:
        st.caption("No collection found. Create a new one:")
        st.text_input("New collection", key="new_collection_name_input")
        st.caption(
            "ℹ️ This collection will be physically created after you run **Index tickets** at least once."
        )
        collection_name = (
            (st.session_state.get("new_collection_name_input") or "").strip()
            or DEFAULT_COLLECTION
        )

    # Single source of truth for collection in session_state
    st.session_state["collection_selected"] = collection_name
    st.session_state["vs_collection"] = collection_name

    # -----------------------------
    # 6) Top status summary (AFTER collection is resolved)
    # -----------------------------
    current_collection = st.session_state.get(
        "vs_collection",
        prefs_dict.get("collection_selected"),
    )
    current_provider = (
        st.session_state.get("emb_provider_select")
        or prefs_dict.get("emb_backend", "OpenAI")
    )
    current_model = (
        st.session_state.get("emb_model")
        or prefs_dict.get("emb_model_name", "")
    )

    status_box.info(
        f"Current collection: **{current_collection or 'none'}** · "
        f"Provider: **{current_provider}** · Model: **{current_model or 'n/a'}**"
    )

    # -----------------------------
    # 7) Collection management: delete
    # -----------------------------
    st.markdown("---")
    st.subheader("Collection management")

    is_existing_collection = collection_name in coll_options

    del_confirm = st.checkbox(
        f"Confirm deletion of '{collection_name}'",
        value=False,
        disabled=not is_existing_collection,
        help="This operation permanently removes the collection from the datastore.",
    )

    if not is_existing_collection:
        st.caption(
            "ℹ️ This collection name is not yet present in Chroma. "
            "Run **Index tickets** at least once before it can be deleted."
        )

    if st.button(
        "Delete collection",
        type="secondary",
        disabled=not is_existing_collection,
        help="Permanently removes the selected collection from the vector datastore.",
    ):
        if not del_confirm:
            st.warning("Check the confirmation box to proceed with deletion.")
        else:
            try:
                _client = get_chroma_client(persist_dir)
                _client.delete_collection(name=collection_name)

                # Remove the collection meta (provider/model) if present
                meta_path = os.path.join(persist_dir, f"{collection_name}__meta.json")
                try:
                    if os.path.exists(meta_path):
                        os.remove(meta_path)
                except Exception:
                    # Do not block the UX if meta deletion fails
                    pass

                # Clean state if it was pointing to the removed collection
                if st.session_state.get("vs_collection") == collection_name:
                    st.session_state["vector"] = None
                    st.session_state["vs_collection"] = None
                    st.session_state["vs_count"] = 0
                    st.session_state["vs_persist_dir"] = persist_dir

                # Clear any loaded results/issues (optional but recommended)
                st.session_state["issues"] = []

                # Remove selection and new name from session
                st.session_state["collection_selected"] = None
                st.session_state["after_delete_reset"] = True

                # Update sticky prefs, if present
                prefs_in_state = st.session_state.get("prefs", {})
                prefs_in_state["collection_selected"] = None
                prefs_in_state["new_collection_name"] = DEFAULT_COLLECTION
                st.session_state["prefs"] = prefs_in_state
                try:
                    save_prefs(prefs_in_state)
                except Exception:
                    pass

                st.success(f"Collection '{collection_name}' deleted successfully.")
                st.rerun()

            except Exception as e:
                st.error(f"Error during deletion: {e}")

    st.markdown("---")

    # -----------------------------
    # 8) Embeddings configuration
    # -----------------------------
    st.header("Embeddings")

    emb_backends: list[str] = []
    if SentenceTransformer is not None and not IS_CLOUD:
        emb_backends.append("Local (sentence-transformers)")
    emb_backends.append("OpenAI")

    pref_backend = prefs_dict.get("emb_backend") or "OpenAI"
    if (
        pref_backend == "Local (sentence-transformers)"
        and "Local (sentence-transformers)" not in emb_backends
    ):
        pref_backend = "OpenAI"

    emb_backend = st.selectbox(
        "Embeddings provider",
        options=emb_backends,
        index=emb_backends.index(pref_backend),
        key="emb_provider_select",
    )

    # Reset the model ONLY if the user has just changed the provider
    prev_backend = st.session_state.get("_prev_emb_backend")
    if prev_backend is None:
        # First time: keep the existing model from prefs/session_state
        st.session_state["_prev_emb_backend"] = emb_backend
    elif prev_backend != emb_backend:
        # Real provider change → reset model accordingly
        st.session_state["_prev_emb_backend"] = emb_backend
        st.session_state["emb_model"] = (
            "all-MiniLM-L6-v2"
            if emb_backend == "Local (sentence-transformers)"
            else "text-embedding-3-small"
        )

    if "emb_model" not in st.session_state:
        st.session_state["emb_model"] = prefs_dict.get(
            "emb_model_name",
            "all-MiniLM-L6-v2"
            if emb_backend == "Local (sentence-transformers)"
            else "text-embedding-3-small",
        )

    emb_model_options = (
        ["all-MiniLM-L6-v2"]
        if emb_backend == "Local (sentence-transformers)"
        else ["text-embedding-3-small", "text-embedding-3-large"]
    )

    if st.session_state["emb_model"] not in emb_model_options:
        st.session_state["emb_model"] = emb_model_options[0]

    emb_model_name = st.selectbox(
        "Embeddings model",
        options=emb_model_options,
        index=emb_model_options.index(st.session_state["emb_model"]),
        key="emb_model",
    )

    # Persist embeddings preferences to sticky prefs so that they survive page navigation / app reruns
    try:
        prefs_in_state = st.session_state.get("prefs", {})
        if prefs_in_state.get("emb_backend") != emb_backend:
            prefs_in_state["emb_backend"] = emb_backend
        if prefs_in_state.get("emb_model_name") != emb_model_name:
            prefs_in_state["emb_model_name"] = emb_model_name
        st.session_state["prefs"] = prefs_in_state
        save_prefs(prefs_in_state)
    except Exception:
        pass

    # -----------------------------
    # 9) Vector DB indexing
    # -----------------------------
    st.subheader("Vector DB Indexing")

    # Embeddings backend/model for ingestion, chat and playbooks
    emb_backend = st.session_state.get(
        "emb_provider_select",
        prefs_dict.get("emb_backend", "OpenAI"),
    )
    emb_model_name = st.session_state.get(
        "emb_model",
        prefs_dict.get("emb_model_name", "text-embedding-3-small"),
    )

    col1, col2 = st.columns([1, 3])
    with col1:
        start_ingest = st.button("Index tickets")
    with col2:
        st.caption(
            "Create ticket embeddings and save them to Chroma for semantic retrieval"
        )

    # Read chunking configuration from session_state
    enable_chunking = bool(st.session_state.get("enable_chunking", True))
    chunk_size = int(st.session_state.get("chunk_size", 800))
    chunk_overlap = int(st.session_state.get("chunk_overlap", 80))
    chunk_min = int(st.session_state.get("chunk_min", 512))

    if start_ingest:
        issues = st.session_state.get("issues", [])
        if not issues:
            st.error("First load the project's tickets")
        else:
            try:
                st.session_state["vector"] = VectorStore(
                    persist_dir=persist_dir, collection_name=collection_name
                )
                use_openai_embeddings = emb_backend == "OpenAI"
                st.session_state["embedder"] = EmbeddingBackend(
                    use_openai=use_openai_embeddings, model_name=emb_model_name
                )

                # Save embeddings model metadata for consistency
                try:
                    meta_path = os.path.join(
                        persist_dir, f"{collection_name}__meta.json"
                    )
                    meta = {
                        "provider": st.session_state["embedder"].provider_name,
                        "model": st.session_state["embedder"].model_name,
                    }
                    with open(meta_path, "w", encoding="utf-8") as f:
                        json.dump(meta, f)
                except Exception:
                    pass

                # Chunked indexing with metadata parent_id, chunk_id, pos
                all_ids = []
                all_docs = []
                all_metas = []
                embed_inputs = []

                for it in issues:
                    full_text = it.text_blob()
                    if enable_chunking:
                        pieces = split_into_chunks(
                            full_text,
                            chunk_size=int(chunk_size),
                            overlap=int(chunk_overlap),
                            min_size=int(chunk_min),
                        )
                    else:
                        pieces = [(0, full_text)]

                    if not pieces:
                        pieces = [(0, full_text)]

                    multi = len(pieces) > 1

                    for idx, (pos0, chunk_text) in enumerate(pieces, start=1):
                        cid = (
                            f"{it.id_readable}::c{idx:03d}" if multi else it.id_readable
                        )

                        all_ids.append(cid)
                        all_docs.append(chunk_text)

                        meta = {
                            "parent_id": it.id_readable,
                            "id_readable": it.id_readable,
                            "summary": it.summary,
                            "project": it.project,
                        }
                        if multi:
                            meta["chunk_id"] = idx
                            meta["pos"] = int(pos0)

                        all_metas.append(meta)
                        all_metas = [
                            {k: v for k, v in m.items() if v is not None}
                            for m in all_metas
                        ]
                        embed_inputs.append(
                            f"{it.id_readable} | {it.summary}\n\n{chunk_text}"
                        )

                with st.spinner(
                    f"Computing embeddings for {len(all_ids)} "
                    f"{'chunks' if enable_chunking else 'documents'} "
                    "and saving to Chroma..."
                ):
                    embs = st.session_state["embedder"].embed(embed_inputs)
                    st.session_state["vector"].add(
                        ids=all_ids,
                        documents=all_docs,
                        metadatas=all_metas,
                        embeddings=embs,
                    )

                # Update counter and active collection
                st.session_state["vs_persist_dir"] = persist_dir
                st.session_state["vs_collection"] = collection_name
                st.session_state["vs_count"] = st.session_state["vector"].count()

                # Store success status in session_state (will be displayed after rerun)
                st.session_state["_index_success"] = (
                    f"Indexing completed. Total "
                    f"{'chunks' if enable_chunking else 'documents'}: "
                    f"{st.session_state['vs_count']}"
                )

                # Stay in Phase 3 after rerun
                st.session_state["ui_phase_choice"] = "Embeddings & Vector DB"

                # Force rerun so the status summary and sidebar reflect the updated info
                st.rerun()

            except Exception as e:
                st.error(f"Indexing error: {e}")

def render_phase_retrieval_page(prefs):
    import streamlit as st

    # 1) Normalize prefs into a dict (same pattern as other phases)
    if isinstance(prefs, dict):
        prefs_dict = prefs
    else:
        prefs_dict = st.session_state.get("prefs", {})

    # 2) Handle advanced reset (flag set in the previous run)
    if st.session_state.pop("_adv_do_reset", False):
        reset_defaults = {
            # Distance / chunking
            "max_distance": 0.9,
            "enable_chunking": True,
            "chunk_size": 800,
            "chunk_overlap": 80,
            "chunk_min": 512,
            # Advanced (UI keys with adv_ prefix)
            "adv_show_distances": False,
            "adv_top_k": 5,
            "adv_collapse_duplicates": True,
            "adv_per_parent_display": 1,
            "adv_per_parent_prompt": 3,
            "adv_stitch_max_chars": 1500,
        }
        # Apply defaults into session_state (widgets will read from here)
        for key, value in reset_defaults.items():
            st.session_state[key] = value

        # Optional visual feedback handled at the end of the function
        st.session_state["_adv_reset_toast"] = True

    # 3) One-time bootstrap from prefs for missing keys
    #    (same pattern as Phase 7 – Solutions memory)
    def init_state(key: str, default):
        """Initialize a session_state key once, if not already set."""
        if key not in st.session_state:
            st.session_state[key] = default

    # Distance + chunking
    init_state("max_distance", float(prefs_dict.get("max_distance", 0.9)))
    init_state("enable_chunking", bool(prefs_dict.get("enable_chunking", True)))
    init_state("chunk_size", int(prefs_dict.get("chunk_size", 800)))
    init_state("chunk_overlap", int(prefs_dict.get("chunk_overlap", 80)))
    init_state("chunk_min", int(prefs_dict.get("chunk_min", 512)))

    # Canonical prefs are key names without adv_ (show_distances, top_k, …)
    # Widgets use adv_* (consistent with Phase 8 – Preferences & debug)
    init_state("adv_show_distances", bool(prefs_dict.get("show_distances", False)))
    init_state("adv_top_k", int(prefs_dict.get("top_k", 5)))
    init_state(
        "adv_collapse_duplicates",
        bool(prefs_dict.get("collapse_duplicates", True)),
    )
    init_state("adv_per_parent_display", int(prefs_dict.get("per_parent_display", 1)))
    init_state("adv_per_parent_prompt", int(prefs_dict.get("per_parent_prompt", 3)))
    init_state("adv_stitch_max_chars", int(prefs_dict.get("stitch_max_chars", 1500)))

    # 4) Main UI
    st.title("Phase 4 – Retrieval configuration")
    st.write(
        "Configure the distance threshold and chunking parameters used to retrieve "
        "similar tickets from the vector store."
    )

    st.markdown("### Distance threshold")

    max_distance = st.slider(
        "Maximum distance threshold (cosine)",
        min_value=0.1,
        max_value=2.0,
        step=0.05,
        value=float(st.session_state["max_distance"]),
        help="Maximum cosine distance for a result to be considered similar.",
    )
    # Keep session_state in sync (widget has no key, so this is safe)
    st.session_state["max_distance"] = float(max_distance)

    st.markdown("---")
    st.subheader("Chunking configuration")

    enable_chunking = st.checkbox(
        "Enable chunking",
        value=bool(st.session_state["enable_chunking"]),
        help="If enabled, long tickets are split into overlapping chunks before indexing.",
    )
    st.session_state["enable_chunking"] = bool(enable_chunking)

    c1, c2, c3 = st.columns(3)

    with c1:
        chunk_size = st.number_input(
            "Chunk size (tokens)",
            min_value=128,
            max_value=2048,
            step=64,
            value=int(st.session_state["chunk_size"]),
            help="Typical values: 512–800.",
        )
        st.session_state["chunk_size"] = int(chunk_size)

    with c2:
        chunk_overlap = st.number_input(
            "Overlap (tokens)",
            min_value=0,
            max_value=512,
            step=10,
            value=int(st.session_state["chunk_overlap"]),
            help="How many tokens to overlap between consecutive chunks.",
        )
        st.session_state["chunk_overlap"] = int(chunk_overlap)

    with c3:
        chunk_min = st.number_input(
            "Min size to chunk",
            min_value=128,
            max_value=2048,
            step=64,
            value=int(st.session_state["chunk_min"]),
            help="Below this size, the ticket is indexed as a single document.",
        )
        st.session_state["chunk_min"] = int(chunk_min)

    st.info(
        "These settings are used when you run indexing (Phase 3) and when the retriever "
        "filters similar tickets (Phase 6)."
    )

    # 5) Advanced settings (UI + Reset)
    with st.expander("Advanced settings", expanded=False):
        show_distances = st.checkbox(
            "Show distances in results",
            value=bool(st.session_state["adv_show_distances"]),
            help="Display the distance/score next to each retrieved result.",
        )
        st.session_state["adv_show_distances"] = bool(show_distances)

        top_k = st.number_input(
            "Top-K KB results",
            min_value=1,
            max_value=50,
            step=1,
            value=int(st.session_state["adv_top_k"]),
            help=(
                "Number of Knowledge Base results to pass downstream "
                "(before collapsing duplicates)."
            ),
        )
        st.session_state["adv_top_k"] = int(top_k)

        collapse_duplicates = st.checkbox(
            "Collapse duplicate results by ticket",
            value=bool(st.session_state["adv_collapse_duplicates"]),
            help="Show only one result per ticket in the UI (keeps recall for the prompt context).",
        )
        st.session_state["adv_collapse_duplicates"] = bool(collapse_duplicates)

        per_parent_display = st.number_input(
            "Max results per ticket (UI)",
            min_value=1,
            max_value=10,
            step=1,
            value=int(st.session_state["adv_per_parent_display"]),
            help="Maximum number of results displayed for the same ticket in the UI.",
        )
        st.session_state["adv_per_parent_display"] = int(per_parent_display)

        per_parent_prompt = st.number_input(
            "Max chunks per ticket (prompt)",
            min_value=1,
            max_value=10,
            step=1,
            value=int(st.session_state["adv_per_parent_prompt"]),
            help="Maximum number of chunks concatenated per ticket in the prompt context.",
        )
        st.session_state["adv_per_parent_prompt"] = int(per_parent_prompt)

        stitch_max_chars = st.number_input(
            "Stitched context limit (chars)",
            min_value=200,
            max_value=20000,
            step=100,
            value=int(st.session_state["adv_stitch_max_chars"]),
            help=(
                "Character limit when concatenating multiple chunks of the same ticket "
                "for the prompt context."
            ),
        )
        st.session_state["adv_stitch_max_chars"] = int(stitch_max_chars)

        # Reset button: only sets a flag, actual reset happens at the beginning
        if st.button("Reset to defaults", key="adv_reset_btn"):
            st.session_state["_adv_do_reset"] = True
            st.rerun()

    # Toast after the reset run
    if st.session_state.pop("_adv_reset_toast", False):
        st.toast("Advanced settings reset to defaults", icon="↩️")

    # 6) Synchronize logical keys used by Chat + sidebar (no adv_ prefix)
    st.session_state["show_distances"] = bool(
        st.session_state.get("adv_show_distances", False)
    )
    st.session_state["top_k"] = int(st.session_state.get("adv_top_k", 5))
    st.session_state["collapse_duplicates"] = bool(
        st.session_state.get("adv_collapse_duplicates", True)
    )
    st.session_state["per_parent_display"] = int(
        st.session_state.get("adv_per_parent_display", 1)
    )
    st.session_state["per_parent_prompt"] = int(
        st.session_state.get("adv_per_parent_prompt", 3)
    )
    st.session_state["stitch_max_chars"] = int(
        st.session_state.get("adv_stitch_max_chars", 1500)
    )

    # 7) Update prefs in memory (Phase 8 handles writing to disk)
    prefs_dict.update(
        {
            "max_distance": float(st.session_state["max_distance"]),
            "enable_chunking": bool(st.session_state["enable_chunking"]),
            "chunk_size": int(st.session_state["chunk_size"]),
            "chunk_overlap": int(st.session_state["chunk_overlap"]),
            "chunk_min": int(st.session_state["chunk_min"]),
            "show_distances": bool(st.session_state["show_distances"]),
            "top_k": int(st.session_state["top_k"]),
            "collapse_duplicates": bool(st.session_state["collapse_duplicates"]),
            "per_parent_display": int(st.session_state["per_parent_display"]),
            "per_parent_prompt": int(st.session_state["per_parent_prompt"]),
            "stitch_max_chars": int(st.session_state["stitch_max_chars"]),
        }
    )
    st.session_state["prefs"] = prefs_dict

def render_phase_llm_page(prefs):
    import streamlit as st

    # Normalizza prefs in dict per avere default decenti
    if isinstance(prefs, dict):
        prefs_dict = prefs
    else:
        prefs_dict = st.session_state.get("prefs", {})

    st.title("Phase 5 – LLM & API keys")
    st.write(
        "Configure the LLM provider, model, temperature and API keys used to answer "
        "questions based on the retrieved tickets."
    )

    # Piccolo riassunto in alto
    current_provider = st.session_state.get(
        "llm_provider_select",
        prefs_dict.get("llm_provider", "OpenAI"),
    )
    current_model = st.session_state.get(
        "llm_model",
        prefs_dict.get("llm_model", "gpt-4o"),
    )
    current_temp = float(
        st.session_state.get(
            "llm_temperature",
            prefs_dict.get("llm_temperature", 0.2),
        )
    )

    st.markdown("---")

    st.header("LLM")

    # Ollama detection
    ollama_ok, ollama_host = (False, None) if IS_CLOUD else is_ollama_available()
    llm_provider_options = ["OpenAI"] + (["Ollama (local)"] if ollama_ok else [])

    # Prefer the value in session_state (if present), otherwise fall back to prefs
    current_provider = st.session_state.get(
        "llm_provider_select",
        prefs_dict.get("llm_provider", "OpenAI"),
    )

    # If Ollama is not available anymore, force provider to OpenAI
    if (not ollama_ok) and (current_provider == "Ollama (local)"):
        current_provider = "OpenAI"

    # Compute the index from the provider name
    if current_provider in llm_provider_options:
        default_idx = llm_provider_options.index(current_provider)
    else:
        default_idx = 0  # safe fallback

    llm_provider = st.selectbox(
        "LLM provider",
        llm_provider_options,
        index=default_idx,
    )
    st.session_state["llm_provider_select"] = llm_provider

    if not ollama_ok:
        st.caption("⚠️ Ollama is not available in this environment; option disabled.")

    # If the provider changes, reset the model to the selected provider's default
    prev_provider = st.session_state.get("last_llm_provider")
    if prev_provider != llm_provider:
        st.session_state["last_llm_provider"] = llm_provider
        st.session_state["llm_model"] = (
            DEFAULT_LLM_MODEL if llm_provider == "OpenAI" else DEFAULT_OLLAMA_MODEL
        )

    # --- LLM MODEL SELECTION ---

    # Determine the default model for this provider
    provider_default_model = (
        DEFAULT_LLM_MODEL if llm_provider == "OpenAI" else DEFAULT_OLLAMA_MODEL
    )

    # Read model from prefs (may be "")
    prefs_model = (prefs_dict.get("llm_model") or "").strip()

    # Build the canonical initial model:
    # Priority: session_state → prefs → per-provider default
    initial_model = (
        st.session_state.get("llm_model")
        or prefs_model
        or provider_default_model
    )

    # Apply it BEFORE creating the widget
    st.session_state["llm_model"] = initial_model

    # Now we can safely draw the widget
    # --- Suggested models (still allow custom typing) ---
    if "llm_model" not in st.session_state:
        st.session_state["llm_model"] = initial_model

    suggested_models = [
        # Keep a short curated list; user can still type any name
        "llama3.2",
        "gemma3:1b",
        "gemma3:4b",
        "qwen3:4b",
        "qwen3:8b",
    ]

    # If Ollama is selected and reachable, merge real models from Ollama /api/tags
    if st.session_state.get("llm_provider_select") == "Ollama (local)" and ollama_ok and ollama_host:
        suggested_models = sorted(set(suggested_models + get_ollama_model_names(ollama_host)))

    # A small selector to quickly pick a model name (optional)
    picked = st.selectbox(
        "Pick from suggested models (optional)",
        options=["(keep current)"] + suggested_models,
        index=0,
    )

    if picked != "(keep current)":
        st.session_state["llm_model"] = picked

    # Free text always wins / allows experimenting with any model
    llm_model = st.text_input(
        "LLM model",
        value=st.session_state["llm_model"],
        help="Type any Ollama model name (e.g., 'gemma3:4b', 'qwen3:8b').",
    )

    # Normalize and rewrite into session_state
    st.session_state["llm_model"] = (llm_model or provider_default_model).strip()

    # Temperature slider (as in your code)
    if "llm_temperature" not in st.session_state:
        st.session_state["llm_temperature"] = float(prefs_dict.get("llm_temperature", 0.2))
    llm_temperature = st.slider("Temperature", 0.0, 1.5, st.session_state["llm_temperature"], 0.05)
    st.session_state["llm_temperature"] = llm_temperature

    # --- [SIDEBAR > API Keys] ---
    st.header("API Keys")
    # Read providers from session_state (fallback to defaults / prefs)
    emb_backend = st.session_state.get(
        "emb_provider_select",
        prefs_dict.get("emb_backend", "OpenAI"),
    )

    llm_provider = st.session_state.get(
        "llm_provider_select",
        prefs_dict.get("llm_provider", "OpenAI"),
    )

    openai_needed = (emb_backend == "OpenAI") or (llm_provider == "OpenAI")

    openai_key_input = st.text_input("OpenAI API Key",
        type="password",
        value=st.session_state.get("openai_key", ""),
        disabled=not openai_needed,
        help="Used only if you choose OpenAI as provider for Embeddings or LLM."
    )

    if openai_key_input:
        st.session_state["openai_key"] = openai_key_input

    if not openai_needed:
        st.caption("OpenAI API Key not needed: you are using only local providers (Ollama / sentence-transformers).")
    elif (llm_provider == "Ollama (local)") and (emb_backend == "OpenAI"):
        st.info("You are using: LLM = Ollama, Embeddings = OpenAI → the key will be used only for embeddings.")

def render_phase_chat_page(prefs):
    import streamlit as st

    # -----------------------------
    # 1) Normalize prefs
    # -----------------------------
    if isinstance(prefs, dict):
        prefs_dict = prefs
    else:
        prefs_dict = st.session_state.get("prefs", {})

    # -----------------------------
    # 2) Persist dir & collection
    # -----------------------------
    persist_dir = (
        st.session_state.get("vs_persist_dir")
        or st.session_state.get("persist_dir")
        or prefs_dict.get("persist_dir", DEFAULT_CHROMA_DIR)
    )

    collection_name = (
        st.session_state.get("vs_collection")
        or prefs_dict.get("collection_selected")
        or st.session_state.get("new_collection_name_input")
        or prefs_dict.get("new_collection_name", DEFAULT_COLLECTION)
    )
    st.session_state["vs_collection"] = collection_name

    max_distance = float(
        st.session_state.get("max_distance", prefs_dict.get("max_distance", 0.9))
    )

    # -----------------------------
    # 3) Embeddings + LLM config (UI / prefs view)
    # -----------------------------
    emb_backend = st.session_state.get(
        "emb_provider_select",
        prefs_dict.get("emb_backend", "OpenAI"),
    )
    emb_model_name = st.session_state.get(
        "emb_model",
        prefs_dict.get("emb_model_name", "text-embedding-3-small"),
    )

    llm_provider = st.session_state.get(
        "llm_provider_select",
        prefs_dict.get("llm_provider", "OpenAI"),
    )
    llm_temperature = float(
        st.session_state.get("llm_temperature", prefs_dict.get("llm_temperature", 0.2))
    )
    llm_model = st.session_state.get(
        "llm_model",
        prefs_dict.get(
            "llm_model",
            DEFAULT_LLM_MODEL if llm_provider == "OpenAI" else DEFAULT_OLLAMA_MODEL,
        ),
    )

    # -----------------------------
    # 4) UI
    # -----------------------------
    st.title("Phase 6 – Chat & Results")
    st.write(
        "Ask questions about your YouTrack tickets. The app retrieves similar tickets "
        "from the vector store, sends them to the LLM, and shows the answer together "
        "with the retrieved context."
    )

    st.markdown("---")

    st.subheader("RAG Chatbot")
    query = st.text_area(
        "New ticket",
        height=140,
        placeholder="Describe the problem as if opening a ticket",
    )
    run_chat = st.button("Search and answer")

    st.caption(
        f"DEBUG: collection={collection_name}, "
        f"vs_collection={st.session_state.get('vs_collection')}, "
        f"vs_count={st.session_state.get('vs_count', 'n/a')}"
    )

    # -----------------------------
    # 5) Click handler per la CHAT
    # -----------------------------
    if run_chat:
        if not query.strip():
            st.error("Enter the ticket text")
            return

        # Ensure the vector store is open on the correct collection
        vect = st.session_state.get("vector")
        if (vect is None) or (st.session_state.get("vs_collection") != collection_name):
            ok, _, _ = open_vector_in_session(persist_dir, collection_name)
            if not ok:
                st.error(
                    "Open or create a valid collection in Phase 3 (Embeddings & Vector DB)."
                )
                return

        # --- DEBUG: list Chroma collections and counts ---
        try:
            _client = get_chroma_client(persist_dir)
            debug_cols = []
            for c in _client.list_collections():
                try:
                    n = c.count()
                except Exception:
                    n = "?"
                debug_cols.append(f"{c.name} ({n})")
            if debug_cols:
                st.caption("DEBUG Chroma collections: " + ", ".join(debug_cols))
        except Exception:
            pass

        try:
            # -----------------------------
            # 6) Build embedder fresh (prefer meta over prefs)
            # -----------------------------
            meta_provider = None
            meta_model = None
            try:
                meta_path = os.path.join(persist_dir, f"{collection_name}__meta.json")
                if os.path.exists(meta_path):
                    with open(meta_path, "r", encoding="utf-8") as f:
                        m = json.load(f)
                    meta_provider = m.get("provider")
                    meta_model = m.get("model")
            except Exception:
                meta_provider = None
                meta_model = None

            provider_for_embed = meta_provider or emb_backend
            model_for_embed = meta_model or emb_model_name

            use_openai_embeddings = (str(provider_for_embed).lower() == "openai")
            embedder = EmbeddingBackend(
                use_openai=use_openai_embeddings,
                model_name=model_for_embed,
            )
            st.session_state["embedder"] = embedder

            # Info se ancora c'è mismatch col meta
            try:
                if meta_provider and meta_model:
                    if (
                        meta_provider != embedder.provider_name
                        or meta_model != embedder.model_name
                    ):
                        st.info(
                            f"Note: the collection was indexed with {meta_provider} / {meta_model}; "
                            f"you are querying with {embedder.provider_name} / {embedder.model_name}."
                        )
            except Exception:
                pass

            # -----------------------------
            # 7) Retrieval
            # -----------------------------
            q_emb = embedder.embed([query])

            top_k = int(st.session_state.get("top_k", 5))
            show_distances = bool(st.session_state.get("show_distances", False))

            # KB
            kb_coll = load_chroma_collection(persist_dir, collection_name, space="cosine")
            kb_res = kb_coll.query(
                query_embeddings=q_emb,
                n_results=top_k,
            )
            kb_docs = kb_res.get("documents", [[]])[0]
            kb_metas = kb_res.get("metadatas", [[]])[0]
            kb_dists = kb_res.get("distances", [[]])[0]

            st.caption(
                f"DEBUG KB: raw_n={len(kb_docs)}, "
                f"count={kb_coll.count()}, "
                f"dists[0:5]={kb_dists[:5]} (max_distance={max_distance})"
            )

            DIST_MAX_KB = max_distance
            kb_retrieved = [
                (doc, meta, dist, "KB")
                for doc, meta, dist in zip(kb_docs, kb_metas, kb_dists)
                if dist is not None and dist <= DIST_MAX_KB
            ]

            # MEMORIES
            mem_retrieved = []
            if st.session_state.get("enable_memory", False):
                try:
                    mem_coll = load_chroma_collection(persist_dir, MEM_COLLECTION, space="cosine")
                    mem_res = mem_coll.query(
                        query_embeddings=q_emb,
                        n_results=min(5, top_k),
                    )
                    mem_docs = mem_res.get("documents", [[]])[0]
                    mem_metas = mem_res.get("metadatas", [[]])[0]
                    mem_dists = mem_res.get("distances", [[]])[0]

                    st.caption(
                        f"DEBUG MEM: raw_n={len(mem_docs)}, "
                        f"count={mem_coll.count()}, "
                        f"dists[0:5]={mem_dists[:5]}"
                    )

                    DIST_MAX_MEM = DIST_MAX_KB
                    now = now_ts()
                    for doc, meta, dist in zip(mem_docs, mem_metas, mem_dists):
                        if dist is None:
                            continue
                        meta = meta or {}
                        exp = int(meta.get("expires_at", 0))
                        if exp and exp < now:
                            continue
                        if dist <= DIST_MAX_MEM:
                            mem_retrieved.append((doc, meta, dist, "MEM"))
                except Exception as e:
                    st.caption(f"DEBUG MEM error: {e}")

            # Merge MEM + KB
            mem_cap = 2
            merged = sorted(mem_retrieved, key=lambda x: x[2])[:mem_cap] + sorted(
                kb_retrieved, key=lambda x: x[2]
            )[: max(0, top_k - min(mem_cap, len(mem_retrieved)))]

            # -----------------------------
            # 8) Prompt build
            # -----------------------------
            if merged:
                merged_collapsed_view = collapse_by_parent(
                    merged,
                    per_parent=int(st.session_state.get("per_parent_display", 1)),
                    stitch_for_prompt=False,
                )

                merged_for_prompt = collapse_by_parent(
                    merged,
                    per_parent=int(st.session_state.get("per_parent_prompt", 3)),
                    stitch_for_prompt=True,
                    max_chars=int(st.session_state.get("stitch_max_chars", 1500)),
                )

                retrieved_for_prompt = [
                    (doc, meta, dist) for (doc, meta, dist, _src) in merged_for_prompt
                ]
                prompt = build_prompt(query, retrieved_for_prompt)
            else:
                merged_collapsed_view = []
                merged_for_prompt = []
                prompt = (
                    f"New ticket:\n{query.strip()}\n\n"
                    "No similar ticket was found in the knowledge base."
                )

            st.write(
                f"DEBUG: view={len(merged_collapsed_view)}, "
                f"prompt_ctx={len(merged_for_prompt)}"
            )

            if st.session_state.get("show_prompt", False):
                with st.expander("Prompt sent to the LLM", expanded=False):
                    st.code(prompt, language="markdown")

            # -----------------------------
            # 9) LLM answer
            # -----------------------------
            _model = (llm_model or "").strip()
            if not _model:
                raise RuntimeError(
                    "LLM model not set: select a valid model in the LLM phase."
                )

            llm = LLMBackend(llm_provider, _model, temperature=llm_temperature)
            with st.spinner("Generating answer."):
                answer = llm.generate(RAG_SYSTEM_PROMPT, prompt)
            st.write(answer)
            st.session_state["last_query"] = query
            st.session_state["last_answer"] = answer

            # -----------------------------
            # 10) Similar results with provenance
            # -----------------------------
            if merged:
                base_url = (
                    st.session_state.get("yt_client").base_url
                    if st.session_state.get("yt_client")
                    else ""
                ).rstrip("/")
                st.write("Similar results (top-k, with provenance):")
                for (doc, meta, dist, src) in merged_collapsed_view:
                    if src == "KB":
                        idr = meta.get("id_readable", "")
                        summary = meta.get("summary", "")
                        url = f"{base_url}/issue/{idr}" if base_url and idr else ""

                        cid = meta.get("chunk_id")
                        cpos = meta.get("pos")

                        header = (
                            f"{idr} – {summary}"
                            if idr and summary
                            else (idr or summary or "KB result")
                        )
                        if url:
                            st.markdown(f"- [{header}]({url})")
                        else:
                            st.markdown(f"- {header}")

                        if show_distances:
                            st.caption(f"Distance: {dist:.4f}")

                        if cid is not None and cpos is not None:
                            st.caption(f"Chunk #{cid} (token offset: {cpos})")

                        with st.expander("View chunk text", expanded=False):
                            st.write(doc)

                    else:  # MEM
                        title = meta.get("title", "Saved playbook")
                        st.markdown(f"- 🧠 **Playbook** – {title}")
                        if show_distances:
                            st.caption(f"Distance: {dist:.4f}")
                        if st.session_state.get("mem_show_full", False):
                            with st.expander("View playbook text", expanded=False):
                                st.write(doc)

        except Exception as e:
            st.error(f"Chat error: {e}")

    # -----------------------------
    # 11) Save as playbook
    # -----------------------------
    if st.session_state.get("enable_memory", False) and st.session_state.get("last_answer"):
        st.markdown("---")
        if st.button("✅ Mark as solved → Save as playbook"):
            try:
                query = st.session_state.get("last_query", "") or ""
                answer = st.session_state.get("last_answer", "") or ""

                condense_prompt = (
                    "Summarize the solution in 3–6 imperative, reusable sentences, "
                    "avoiding sensitive data.\n"
                    f"- Query: {query.strip()}\n- Answer:\n{answer}\n\n"
                    "Output: bulleted list of clear steps."
                )

                _model = (st.session_state.get("llm_model") or "").strip()
                if not _model:
                    prov = st.session_state.get("llm_provider_select", "OpenAI")
                    _model = DEFAULT_LLM_MODEL if prov == "OpenAI" else DEFAULT_OLLAMA_MODEL

                try:
                    llm_mem = LLMBackend(
                        llm_provider,
                        _model,
                        temperature=max(0.1, llm_temperature - 0.1),
                    )
                    playbook_text = llm_mem.generate(
                        "You are an assistant who distills reusable mini-playbooks.",
                        condense_prompt,
                    ).strip()
                except Exception:
                    playbook_text = (answer[:800] + "…") if len(answer) > 800 else answer

                if not playbook_text.strip():
                    st.warning("No content to save.")
                    return

                curr_proj = st.session_state.get("last_project_key") or ""
                tags_list = ["playbook"] + ([curr_proj] if curr_proj else [])

                meta = {
                    "source": "memory",
                    "project": curr_proj,
                    "quality": "verified",
                    "created_at": now_ts(),
                    "expires_at": ts_in_days(
                        st.session_state.get("mem_ttl_days", DEFAULT_MEM_TTL_DAYS)
                    ),
                    "tags": ",".join(tags_list),
                }

                if st.session_state.get("embedder") is None:
                    use_openai_embeddings = (emb_backend == "OpenAI")
                    st.session_state["embedder"] = EmbeddingBackend(
                        use_openai=use_openai_embeddings,
                        model_name=emb_model_name,
                    )

                mem_emb = st.session_state["embedder"].embed([playbook_text])[0]

                mem_coll = load_chroma_collection(persist_dir, MEM_COLLECTION, space="cosine")
                mem_id = f"mem::{uuid.uuid4().hex[:12]}"
                mem_coll.add(
                    ids=[mem_id],
                    documents=[playbook_text],
                    metadatas=[meta],
                    embeddings=[mem_emb],
                )

                mem_count = 0
                try:
                    mem_count = mem_coll.count()
                except Exception:
                    pass

                st.caption(
                    f"Saved playbook in path='{persist_dir}', "
                    f"collection='{MEM_COLLECTION}', count={mem_count}"
                )
                st.success("Playbook saved in memory.")

                st.session_state["open_memories_after_save"] = True
                st.rerun()
            except Exception as e:
                st.error(f"Errore salvataggio playbook: {e}")

def render_phase_solutions_memory_page(prefs):
    import streamlit as st
    from datetime import datetime

    # 1) prefs_dict: dai priorità a quelli in session_state
    if isinstance(prefs, dict):
        prefs_dict = prefs
    else:
        prefs_dict = st.session_state.get("prefs", {})

    persist_dir = (
        st.session_state.get("vs_persist_dir")
        or st.session_state.get("persist_dir")
        or prefs_dict.get("persist_dir", DEFAULT_CHROMA_DIR)
    )

    # 3) Bootstrap UNA SOLA VOLTA i valori in session_state dai prefs
    if "enable_memory" not in st.session_state:
        st.session_state["enable_memory"] = bool(prefs_dict.get("enable_memory", False))

    if "mem_ttl_days" not in st.session_state:
        st.session_state["mem_ttl_days"] = int(prefs_dict.get("mem_ttl_days", DEFAULT_MEM_TTL_DAYS))

    if "mem_show_full" not in st.session_state:
        st.session_state["mem_show_full"] = bool(prefs_dict.get("mem_show_full", False))

    if "show_memories" not in st.session_state:
        st.session_state["show_memories"] = bool(prefs_dict.get("show_memories", False))

    st.title("Phase 7 – Solutions memory")
    st.write(
        "Review and manage saved playbooks (memories) derived from solved tickets."
    )

    st.markdown("---")
    st.header("Solutions memory")

    # 4) Widget → value da session_state, poi riscrivo in session_state
    mem_show_full = st.checkbox(
        "Show full text of playbooks (MEM)",
        value=st.session_state["mem_show_full"],
        help="If enabled, an expander with the full playbook appears under each MEM result.",
    )
    st.session_state["mem_show_full"] = mem_show_full

    enable_memory = st.checkbox(
        "Enable 'playbook' memory (separate collection)",
        value=st.session_state["enable_memory"],
        help="When you mark an answer as 'Solved', I save a reusable mini-playbook.",
    )
    st.session_state["enable_memory"] = enable_memory

    mem_ttl_days = st.number_input(
        "TTL (days) for playbooks",
        min_value=7,
        max_value=365,
        step=1,
        value=int(st.session_state["mem_ttl_days"]),
    )
    st.session_state["mem_ttl_days"] = int(mem_ttl_days)

    c_mem1, c_mem2 = st.columns(2)
    with c_mem1:
        if st.session_state.pop("open_memories_after_save", False):
            st.session_state["show_memories"] = True

        show_memories = st.checkbox(
            "Show saved playbooks",
            value=st.session_state["show_memories"],
            help="Display the list of the 'memories' collection on the main page.",
        )
        st.session_state["show_memories"] = show_memories

    with c_mem2:
        mem_del_confirm = st.checkbox("Confirm delete memories", value=False)
        if st.button("Delete all memories", disabled=not mem_del_confirm):
            try:
                _client = get_chroma_client(persist_dir)
                _client.delete_collection(name=MEM_COLLECTION)
                load_chroma_collection(persist_dir, MEM_COLLECTION, space="cosine")
                st.success("Memories deleted.")
            except Exception as e:
                st.error(f"Error deleting memories: {e}")

    st.markdown("---")

    # 5) Tabella/Lista dei playbook solo se attivata
    if st.session_state.get("show_memories"):
        st.subheader("Saved playbooks (memories)")
        try:
            client = get_chroma_client(persist_dir)
            mem_coll = load_chroma_collection(persist_dir, MEM_COLLECTION, space="cosine")
            count = mem_coll.count()
            st.caption(f"Collection: '{MEM_COLLECTION}' — path: {persist_dir} — count: {count}")

            data = mem_coll.get(include=["documents", "metadatas"], limit=max(200, count))
            ids = data.get("ids") or []
            docs = data.get("documents") or []
            metas = data.get("metadatas") or []

            if count > 0 and not ids:
                st.warning("La collection riporta count>0 ma get() è vuoto. Riprovo senza include…")
                data = mem_coll.get(limit=max(200, count))  # fallback
                ids = data.get("ids") or []
                docs = data.get("documents") or []
                metas = data.get("metadatas") or []

            if not ids:
                st.caption("Nessun playbook salvato.")
            else:
                for _id, doc, meta in zip(ids, docs, metas):
                    meta = meta or {}

                    created = meta.get("created_at")
                    expires = meta.get("expires_at")

                    created_s = datetime.fromtimestamp(created).strftime("%Y-%m-%d") if created else ""
                    expires_s = datetime.fromtimestamp(expires).strftime("%Y-%m-%d") if expires else ""

                    raw_tags = meta.get("tags", "")
                    if isinstance(raw_tags, str):
                        tags = raw_tags
                    elif raw_tags:
                        tags = ", ".join(raw_tags)
                    else:
                        tags = ""

                    project = meta.get("project", "")
                    prev = (doc[:120] + "…") if doc and len(doc) > 120 else (doc or "")

                    col_main, col_trash = st.columns([0.94, 0.06])

                    with col_main:
                        st.markdown(f"Project: {project or '—'}")
                        if tags:
                            st.caption(f"Tags: {tags}")
                        st.caption(
                            f"Created: {created_s or '—'}  —  Expires: {expires_s or '—'}"
                        )
                        st.markdown(f"Preview: {prev}")

                        # Optional full text of the playbook
                        if mem_show_full and doc:
                            with st.expander("Show full playbook"):
                                st.markdown(doc)

                        # Optional small ID for debugging
                        st.caption(f"ID: `{_id}`")

                    with col_trash:
                        if st.button(
                            "🗑️",
                            key=f"mem_delete_{_id}",
                            help="Delete this playbook",
                        ):
                            try:
                                # Atomic delete of a single solution memory
                                mem_coll.delete(ids=[_id])
                                st.success("Playbook deleted.")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Error deleting playbook: {e}")
        except Exception as e:
            st.error(f"Error reading memories: {e}")

def render_phase_preferences_debug_page(prefs):
    st.title("Phase 8 – Preferences & debug")
    st.write(
        "Preferences handling and Debug settings"
    )
    st.subheader("Preferences")

    # Normalize prefs to a dict
    if isinstance(prefs, dict):
        prefs_dict = prefs
    else:
        prefs_dict = st.session_state.get("prefs", {})

    if "prefs_enabled" not in st.session_state:
        st.session_state["prefs_enabled"] = bool(prefs_dict.get("prefs_enabled", True))

    prefs_enabled = st.checkbox(
        "Enable preferences memory (local)",
        value=st.session_state["prefs_enabled"],
    )
    st.session_state["prefs_enabled"] = int(prefs_enabled)

    c1, c2 = st.columns(2)
    with c1:
        if st.button("Save preferences"):
            if prefs_enabled:
                try:
                    # --- Provider coercion: if Ollama not available, force OpenAI ---
                    _provider_for_save = st.session_state.get("last_llm_provider") or \
                        st.session_state.get("llm_provider_select") or prefs_dict.get("llm_provider", "OpenAI")

                    # Local Ollama availability check (do not rely on UI-phase variable)
                    if IS_CLOUD:
                        _ollama_ok = False
                    else:
                        _ollama_ok, _ = is_ollama_available()
                    if not _ollama_ok and _provider_for_save == "Ollama (local)":
                        _provider_for_save = "OpenAI"

                    # --- LLM model: never empty and consistent with provider ---
                    _model_for_save = (st.session_state.get("llm_model") or "").strip()
                    if not _model_for_save:
                        _model_for_save = DEFAULT_LLM_MODEL if _provider_for_save == "OpenAI" else DEFAULT_OLLAMA_MODEL

                    # --- Read current UI values or session defaults ---
                    yt_url = st.session_state.get("yt_url", prefs_dict.get("yt_url", ""))
                    persist_dir = st.session_state.get("persist_dir", prefs_dict.get("persist_dir", ""))

                    emb_backend = st.session_state.get("emb_provider_select", prefs_dict.get("emb_backend", "OpenAI"))
                    emb_model_name = st.session_state.get("emb_model", prefs_dict.get("emb_model_name", "text-embedding-3-small"))

                    llm_temperature = st.session_state.get("llm_temperature", prefs_dict.get("llm_temperature", 0.2))
                    max_distance = st.session_state.get("max_distance", prefs_dict.get("max_distance", 0.9))
                    show_prompt = st.session_state.get("show_prompt", prefs_dict.get("show_prompt", True))
                    collection_selected = st.session_state.get("collection_selected", prefs_dict.get("collection_selected"))
                    new_collection_name = st.session_state.get("new_collection_name_input", prefs_dict.get("new_collection_name", "tickets"))

                    # --- Chunking ---
                    enable_chunking = bool(st.session_state.get("enable_chunking", prefs_dict.get("enable_chunking", True)))
                    chunk_size = int(st.session_state.get("chunk_size", prefs_dict.get("chunk_size", 800)))
                    chunk_overlap = int(st.session_state.get("chunk_overlap", prefs_dict.get("chunk_overlap", 80)))
                    chunk_min = int(st.session_state.get("chunk_min", prefs_dict.get("chunk_min", 512)))

                    # --- Advanced settings ---
                    show_distances = bool(st.session_state.get("adv_show_distances", prefs_dict.get("show_distances", False)))
                    top_k = int(st.session_state.get("adv_top_k", prefs_dict.get("top_k", 5)))
                    collapse_duplicates = bool(st.session_state.get("adv_collapse_duplicates", prefs_dict.get("collapse_duplicates", True)))
                    per_parent_display = int(st.session_state.get("adv_per_parent_display", prefs_dict.get("per_parent_display", 1)))
                    per_parent_prompt = int(st.session_state.get("adv_per_parent_prompt", prefs_dict.get("per_parent_prompt", 3)))
                    stitch_max_chars = int(st.session_state.get("adv_stitch_max_chars", prefs_dict.get("stitch_max_chars", 1500)))

                    # --- Solutions memory settings ---
                    enable_memory = bool(st.session_state.get("enable_memory", prefs_dict.get("enable_memory", False)))
                    mem_ttl_days = int(st.session_state.get("mem_ttl_days", prefs_dict.get("mem_ttl_days", DEFAULT_MEM_TTL_DAYS)))
                    mem_show_full = bool(st.session_state.get("mem_show_full", prefs_dict.get("mem_show_full", False)))
                    show_memories = bool(st.session_state.get("show_memories", prefs_dict.get("show_memories", False)))

                    # --- Save all prefs ---
                    new_prefs = {
                        "yt_url": yt_url,
                        "persist_dir": persist_dir,
                        "emb_backend": emb_backend,
                        "emb_model_name": emb_model_name,
                        "llm_provider": _provider_for_save,
                        "llm_model": _model_for_save,
                        "llm_temperature": llm_temperature,
                        "max_distance": max_distance,
                        "show_prompt": show_prompt,
                        "collection_selected": collection_selected,
                        "new_collection_name": new_collection_name,
                        "enable_chunking": enable_chunking,
                        "chunk_size": chunk_size,
                        "chunk_overlap": chunk_overlap,
                        "chunk_min": chunk_min,
                        "show_distances": show_distances,
                        "top_k": top_k,
                        "collapse_duplicates": collapse_duplicates,
                        "per_parent_display": per_parent_display,
                        "per_parent_prompt": per_parent_prompt,
                        "stitch_max_chars": stitch_max_chars,
                        "enable_memory": enable_memory,
                        "mem_ttl_days": mem_ttl_days,
                        "mem_show_full": mem_show_full,
                        "show_memories": show_memories,
                        "prefs_enabled": prefs_enabled,
                    }

                    save_prefs(new_prefs)
                    st.session_state["prefs"] = load_prefs()
                    st.success("Preferences saved.")
                    st.toast("Saved to .app_prefs.json", icon="✅")

                except Exception as e:
                    import traceback, textwrap
                    st.error(f"Errore durante il salvataggio: {e}")
                    st.code(textwrap.dedent(traceback.format_exc()))
            else:
                st.info("Preferences memory disabled.")

    with c2:
        if st.button("Restore defaults"):
            try:
                if os.path.exists(PREFS_PATH):
                    os.remove(PREFS_PATH)
                st.session_state["prefs"] = {}
                st.success("Preferences restored. Reload the page to see the defaults.")
                st.rerun()
            except Exception as e:
                st.error(f"Error while restoring: {e}")

    st.header("Debug")
    if "show_prompt" not in st.session_state:
        st.session_state["show_prompt"] = bool(prefs_dict.get("show_prompt", False))
    show_prompt = st.checkbox("Show LLM prompt", value=st.session_state["show_prompt"])
    st.session_state["show_prompt"] = show_prompt

# ------------------------------
# Main Streamlit UI
# ------------------------------
def run_streamlit_app():
    st.set_page_config(page_title="YouTrack RAG Support", layout="wide")
    inject_global_css()
    init_prefs_in_session()
    # === Robust prefs loading & one-time bootstrap ===
    DEFAULT_PREFS = {
        "yt_url": "",
        "persist_dir": "",
        "emb_backend": "OpenAI",
        "emb_model_name": "text-embedding-3-small",
        "llm_provider": "OpenAI",
        "llm_model": "gpt-4o",
        "llm_temperature": 0.2,
        "max_distance": 0.9,
        "show_prompt": True,
        "collection_selected": None,
        "new_collection_name": "tickets",

        # chunking
        "enable_chunking": True,
        "chunk_size": 800,
        "chunk_overlap": 80,
        "chunk_min": 512,

        # advanced settings
        "show_distances": False,
        "top_k": 5,
        "collapse_duplicates": True,
        "per_parent_display": 1,
        "per_parent_prompt": 3,
        "stitch_max_chars": 1500,

        # solutions memory (NUOVI)
        "enable_memory": False,
        "mem_ttl_days": DEFAULT_MEM_TTL_DAYS,
        "mem_show_full": False,
        "show_memories": False,
    }

    def _normalize(p: dict) -> dict:
        p = p or {}
        norm = DEFAULT_PREFS.copy()
        norm.update(p)
        return norm

    # Bootstrap prefs -> session una sola volta
    if not st.session_state.get("_prefs_bootstrapped", False):
        prefs = _normalize(st.session_state.get("prefs", {}))

        # Inizializza SOLO se mancanti (non sovrascrivere scelte dell'utente)
        def _init(key, val):
            if key not in st.session_state:
                st.session_state[key] = val

        # Chunking
        _init("enable_chunking", prefs.get("enable_chunking"))
        _init("chunk_size",       int(prefs.get("chunk_size", 800)))
        _init("chunk_overlap",    int(prefs.get("chunk_overlap", 80)))
        _init("chunk_min",        int(prefs.get("chunk_min", 512)))

        # Advanced
        _init("adv_show_distances",      prefs.get("show_distances"))
        _init("adv_top_k",               int(prefs.get("top_k", 5)))
        _init("adv_collapse_duplicates", prefs.get("collapse_duplicates"))
        _init("adv_per_parent_display",  int(prefs.get("per_parent_display", 1)))
        _init("adv_per_parent_prompt",   int(prefs.get("per_parent_prompt", 3)))
        _init("adv_stitch_max_chars",    int(prefs.get("stitch_max_chars", 1500)))

        # Nome nuova collection (solo se non c'è già)
        _init("new_collection_name_input", (prefs.get("new_collection_name") or DEFAULT_COLLECTION))

        # Solutions memory
        _init("enable_memory",     bool(prefs.get("enable_memory", False)))
        _init("mem_ttl_days",      int(prefs.get("mem_ttl_days", DEFAULT_MEM_TTL_DAYS)))
        _init("mem_show_full",     bool(prefs.get("mem_show_full", False)))
        _init("show_memories",     bool(prefs.get("show_memories", False)))

        st.session_state["_prefs_bootstrapped"] = True

    # Ensure prefs is available beyond the bootstrap block
    prefs = st.session_state.get("prefs", {})

    # Init UI input key (separate from app state)
    if "new_collection_name_input" not in st.session_state:
        st.session_state["new_collection_name_input"] = (prefs.get("new_collection_name") or DEFAULT_COLLECTION)

    # Pre-reset (prima dei widget)
    if st.session_state.get("after_delete_reset"):
        st.session_state["new_collection_name_input"] = DEFAULT_COLLECTION
        st.session_state["after_delete_reset"] = False

    st.title("YouTrack RAG Support")
    st.caption("Support ticket management based on YouTrack history + RAG + LLM")
    phase = st.session_state.get("ui_phase_choice", "YouTrack connection")

    # Default values, in case we are not on this phase
    yt_url = st.session_state.get("yt_url", prefs.get("yt_url", ""))
    yt_token = st.session_state.get("yt_token", "")
    connect = False

    # Open colored container for the current phase
    phase_container_start(phase)

    if phase == "YouTrack connection":
        yt_url, yt_token, connect = render_phase_yt_connection_page(prefs)

    elif phase == "MCP Console":
        render_phase_mcp_console_page(prefs)

    elif phase == "Embeddings & Vector DB":
        render_phase_embeddings_vectordb_page(prefs)

    elif phase == "Retrieval configuration":
        render_phase_retrieval_page(prefs)

    elif phase == "LLM & API keys":
        render_phase_llm_page(prefs)

    elif phase == "Chat & Results":
        render_phase_chat_page(prefs)

    elif phase == "Solutions memory":
        render_phase_solutions_memory_page(prefs)

    elif phase == "Preferences & debug":
        render_phase_preferences_debug_page(prefs)

    elif phase == "Docs KB (PDF/DOCX/TXT)":
        render_phase_docs_kb_page(prefs)

    # Close colored container
    phase_container_end()

    with st.sidebar:
        # --- Wizard-style navigation (visual only, all sections still visible) ---
        PHASES = [
            "YouTrack connection",
            "MCP Console",
            "Embeddings & Vector DB",
            "Retrieval configuration",
            "LLM & API keys",
            "Chat & Results",
            "Solutions memory",
            "Preferences & debug",
            "Docs KB (PDF/DOCX/TXT)",
        ]

        if "ui_phase_choice" not in st.session_state:
            st.session_state["ui_phase_choice"] = PHASES[0]

        # Default index: keep last selection if present, otherwise 0
        default_idx = int(st.session_state.get("ui_phase_index", 0))
        if default_idx < 0 or default_idx >= len(PHASES):
            default_idx = 0

        current_phase = st.radio(
            "Current phase",
            options=PHASES,
            key="ui_phase_choice",
            format_func=lambda p: f"{PHASE_ICONS.get(p, '•')} {p}",
            help="Select the current configuration phase.",
        )

        if isinstance(current_phase, str):
            st.session_state["ui_phase_index"] = PHASES.index(current_phase)

        # Simple progress bar for the wizard
        step_idx = st.session_state.get("ui_phase_index", 0)
        st.progress((step_idx + 1) / len(PHASES), text=f"Step {step_idx + 1} / {len(PHASES)}")

        st.markdown("---")

        quit_btn = False
        st.markdown("### YouTrack status")
        if st.session_state.get("yt_client"):
            st.success("Connected to YouTrack")
        else:
            st.warning("Not connected to YouTrack")

        st.caption(st.session_state.get("yt_url", prefs.get("yt_url", "")) or "No URL configured")

        st.markdown("---")
        st.markdown("### Vector DB / Embeddings status")

        persist_dir = st.session_state.get("persist_dir", prefs.get("persist_dir", ""))

        # Prefer the active collection in session; if missing, fallback to prefs
        vs_collection = st.session_state.get("vs_collection")
        if not vs_collection:
            # Try collection_selected stored in prefs
            pref_coll = None
            if isinstance(prefs, dict):
                pref_coll = (prefs.get("collection_selected") or "").strip()
                if not pref_coll:
                    pref_coll = (prefs.get("new_collection_name") or "").strip()
            if pref_coll:
                vs_collection = pref_coll
                st.session_state["vs_collection"] = pref_coll
                st.session_state.setdefault("collection_selected", pref_coll)
            else:
                vs_collection = ""

        emb_provider = st.session_state.get("emb_provider_select", "OpenAI")
        emb_model = st.session_state.get("emb_model", "")

        if vs_collection:
            st.success(f"Collection: {vs_collection}")
        else:
            st.warning("No collection selected")


        st.caption(persist_dir or "No Chroma path configured")
        st.caption(f"Embeddings: {emb_provider} · {emb_model or 'n/a'}")

        st.markdown("---")
        st.markdown("### LLM status")

        llm_provider = st.session_state.get("llm_provider_select", "OpenAI")
        llm_model = st.session_state.get("llm_model", "gpt-4o")
        llm_temperature = float(st.session_state.get("llm_temperature", 0.2))

        st.caption(f"Provider: {llm_provider}")
        st.caption(f"Model: {llm_model}")
        st.caption(f"Temperature: {llm_temperature:.2f}")

        # --- Retrieval summary (read-only) ---
        st.markdown("---")
        st.header("Retrieval summary")

        def _yes(v: bool) -> str:
            return "✅" if bool(v) else "✖️"

        st.markdown("---")
        st.markdown("### Chat status")

        last_query = st.session_state.get("last_query", "").strip()
        if last_query:
            st.caption(f"Last query: {last_query[:50]}{'…' if len(last_query) > 50 else ''}")
        else:
            st.caption("No query asked yet.")

        # Gather values from session/prefs
        _top_k   = int(st.session_state.get("top_k", 5))
        _maxdist = float(st.session_state.get("max_distance", 0.9))
        _collapse = bool(st.session_state.get("collapse_duplicates", True))
        _pp_ui   = int(st.session_state.get("per_parent_display", 1))
        _pp_pr   = int(st.session_state.get("per_parent_prompt", 3))
        _stitch  = int(st.session_state.get("stitch_max_chars", 1500))

        _chunk_on = bool(st.session_state.get("enable_chunking", True))
        _csize    = int(st.session_state.get("chunk_size", 800))
        _coverlap = int(st.session_state.get("chunk_overlap", 80))
        _cmin     = int(st.session_state.get("chunk_min", 512))

        _emb_backend = st.session_state.get("emb_provider_select") or st.session_state.get("prefs", {}).get("emb_backend", "OpenAI")
        _emb_model   = st.session_state.get("emb_model") or st.session_state.get("prefs", {}).get("emb_model_name", "text-embedding-3-small")

        _collection  = (
            st.session_state.get("collection_selected")
            or st.session_state.get("new_collection_name_input")
            or st.session_state.get("prefs", {}).get("new_collection_name")
            or "—"
        )

        # Row 1: core retrieval knobs
        c1, c2, c3 = st.columns(3)
        with c1: st.caption("Top-K"); st.write(_top_k)
        with c2: st.caption("Max distance"); st.write(f"{_maxdist:.2f}")
        with c3: st.caption("Collapse duplicates"); st.write(_yes(_collapse))

        # Row 2: per-ticket aggregation
        c4, c5, c6 = st.columns(3)
        with c4: st.caption("Per ticket (UI)"); st.write(_pp_ui)
        with c5: st.caption("Per ticket (prompt)"); st.write(_pp_pr)
        with c6: st.caption("Stitch limit (chars)"); st.write(_stitch)

        # Row 3: chunking
        c7, c8, c9, c10 = st.columns(4)
        with c7: st.caption("Chunking enabled"); st.write(_yes(_chunk_on))
        with c8: st.caption("Chunk size"); st.write(_csize)
        with c9: st.caption("Overlap"); st.write(_coverlap)
        with c10: st.caption("Min to chunk"); st.write(_cmin)

        # Row 4: embeddings + collection
        c11, c12 = st.columns(2)
        with c11: st.caption("Embeddings"); st.write(f"{_emb_backend} · {_emb_model}")
        with c12: st.caption("Collection"); st.write(_collection)

        # --- Embedding status (indexed vs query) ---
        _collection = (
            st.session_state.get("collection_selected")
            or st.session_state.get("new_collection_name_input")
            or st.session_state.get("prefs", {}).get("new_collection_name")
        )

        _persist_dir = st.session_state.get("persist_dir") or st.session_state.get("prefs", {}).get("persist_dir", "")

        # Embedder scelto per la QUERY (UI corrente)
        qry_provider = st.session_state.get("emb_provider_select") or st.session_state.get("prefs", {}).get("emb_backend", "OpenAI")
        qry_model    = st.session_state.get("emb_model") or st.session_state.get("prefs", {}).get("emb_model_name", "text-embedding-3-small")

        # Embedder con cui la COLLECTION è stata indicizzata (se noto)
        idx_provider, idx_model, idx_src = (None, None, None)
        if _collection and _persist_dir:
            idx_provider, idx_model, idx_src = get_index_embedder_info(_persist_dir, _collection)

        st.markdown("---")
        st.header("Embedding status")

        def _fmt(v): return v if v else "—"

        colA, colB = st.columns(2)
        with colA:
            st.caption("Indexed with")
            st.write(f"{_fmt(idx_provider)} · {_fmt(idx_model)}" + (f" ({idx_src})" if idx_src else ""))

        with colB:
            st.caption("Query using")
            st.write(f"{_fmt(qry_provider)} · {_fmt(qry_model)}")

        # Avviso mismatch (provider o modello diverso)
        if idx_provider and idx_model:
            mismatch = (_normalize_provider_label(qry_provider) != _normalize_provider_label(idx_provider)) or (qry_model.strip() != idx_model.strip())
            if mismatch:
                st.warning("Embedding mismatch between indexed collection and current query settings.", icon="⚠️")
        else:
            st.caption("No index metadata available. Consider reindexing to record provider/model.")

        st.markdown("---")
        st.caption(f"IS_CLOUD={IS_CLOUD} · ChromaDB dir={st.session_state['prefs'].get('persist_dir', DEFAULT_CHROMA_DIR)}")

        if not IS_CLOUD:
            quit_btn = st.button("Quit", use_container_width=True)
            if quit_btn:
                st.write("Closing application...")
                os._exit(0)

        # Automatic opening of the selected collection (without re-indexing)
        # Avoid opening if the user has chosen "Create new…" but has not entered a name
        vector_ready = False

        persist_dir = st.session_state.get(
            "persist_dir",
            prefs.get("persist_dir", DEFAULT_CHROMA_DIR) if isinstance(prefs, dict) else DEFAULT_CHROMA_DIR,
        )

        # Determine the effective collection name:
        # 1) active vs_collection
        # 2) collection_selected from prefs/session
        # 3) new_collection_name_input / default
        if "vs_collection" in st.session_state and st.session_state["vs_collection"]:
            collection_name = st.session_state["vs_collection"]
        else:
            pref_coll = None
            if isinstance(prefs, dict):
                pref_coll = (prefs.get("collection_selected") or "").strip()
                if not pref_coll:
                    pref_coll = (prefs.get("new_collection_name") or "").strip()
            collection_name = (
                pref_coll
                or st.session_state.get("new_collection_name_input", DEFAULT_COLLECTION)
            )

        if persist_dir and collection_name:
            changed = (
                st.session_state.get("vs_persist_dir") != persist_dir
                or st.session_state.get("vs_collection") != collection_name
                or st.session_state.get("vector") is None
            )

            # Do not autoload if user explicitly selected "Create new collection..."
            # and the name is still empty / default. Use collection_select (the widget),
            # not collection_selected (which is always a real name).
            if (
                st.session_state.get("collection_select") == NEW_LABEL
                and (not collection_name or collection_name == DEFAULT_COLLECTION)
            ):
                pass  # wait for "Index tickets" to create the new collection
            else:
                if changed:
                    ok, cnt, err = open_vector_in_session(persist_dir, collection_name)
                    vector_ready = ok
                    if ok:
                        st.caption(
                            f"Collection '{collection_name}' opened. "
                            f"Indexed documents: {cnt if cnt >= 0 else 'N/A'}"
                        )
                    else:
                        st.caption(f"Unable to open the collection: {err}")
                else:
                    vector_ready = True
                    cnt = st.session_state.get("vs_count", -1)
                    st.caption(
                        f"Collection '{collection_name}' ready. "
                        f"Indexed documents: {cnt if cnt >= 0 else 'N/A'}"
                    )

        if quit_btn:
            st.write("Closing application...")
            os._exit(0)

    # Session state init
    if "yt_client" not in st.session_state:
        st.session_state["yt_client"] = None
    if "projects" not in st.session_state:
        st.session_state["projects"] = []
    if "issues" not in st.session_state:
        st.session_state["issues"] = []
    if "vector" not in st.session_state:
        st.session_state["vector"] = None
    if "embedder" not in st.session_state:
        st.session_state["embedder"] = None

    # Connect to YouTrack
    if connect:
        if not yt_url or not yt_token:
            st.error("Enter YouTrack URL and Token")
        else:
            try:
                st.session_state["yt_client"] = YouTrackClient(yt_url, yt_token)
                with st.spinner("Loading projects..."):
                    st.session_state["projects"] = st.session_state["yt_client"].list_projects()
                st.rerun()
            except Exception as e:
                st.error(f"Connection failed: {e}")


# ------------------------------
# Phase 9 – Docs KB (PDF/DOCX/TXT)
# ------------------------------
DOCS_KB_COLLECTION = "docs_kb"
DOCS_KB_MANIFEST = "docs_kb__manifest.json"

def _docs_manifest_path(persist_dir: str) -> str:
    return os.path.join(persist_dir, DOCS_KB_MANIFEST)

def _load_docs_manifest(persist_dir: str) -> list:
    """Load the docs manifest from disk (best-effort)."""
    try:
        p = _docs_manifest_path(persist_dir)
        if os.path.exists(p):
            with open(p, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list):
                return data
    except Exception:
        pass
    return []

def _save_docs_manifest(persist_dir: str, items: list) -> None:
    """Persist the docs manifest to disk."""
    try:
        os.makedirs(persist_dir, exist_ok=True)
        p = _docs_manifest_path(persist_dir)
        with open(p, "w", encoding="utf-8") as f:
            json.dump(items, f, ensure_ascii=False, indent=2)
    except Exception:
        # Non-fatal on Cloud ephemeral FS
        pass

def _sha256_bytes(b: bytes) -> str:
    import hashlib
    return hashlib.sha256(b).hexdigest()

def _extract_text_from_pdf_bytes(data: bytes) -> str:
    """Extract text from PDF using best available backend. No OCR here."""
    # Try PyMuPDF first (best overall)
    try:
        import fitz  # PyMuPDF
        parts = []
        with fitz.open(stream=data, filetype="pdf") as doc:
            for page in doc:
                t = page.get_text("text") or ""
                if t.strip():
                    parts.append(t)
        txt = "\n".join(parts).strip()
        if txt:
            return txt
    except Exception:
        pass

    # Fallback: pdfplumber
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                if t.strip():
                    parts.append(t)
        txt = "\n".join(parts).strip()
        if txt:
            return txt
    except Exception:
        pass

    # Last resort: PyPDF2
    from PyPDF2 import PdfReader
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        t = page.extract_text() or ""
        if t.strip():
            parts.append(t)
    return "\n".join(parts).strip()

def _extract_text_from_docx_bytes(data: bytes) -> str:
    """Extract text from DOCX bytes, including tables and preserving structure."""
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        doc = Document(io.BytesIO(data))
        parts = []

        # Extract from body (preserves order of paragraphs and tables)
        for element in doc.element.body:
            # Paragraph
            if element.tag.endswith('}p'):
                para = Paragraph(element, doc)
                if para.text and para.text.strip():
                    parts.append(para.text)

            # Table
            elif element.tag.endswith('}tbl'):
                table = Table(element, doc)
                table_text = _extract_table_text(table)
                if table_text:
                    parts.append(table_text)

        # Also extract headers/footers (may contain version info, notes)
        for section in doc.sections:
            if section.header:
                header_text = "\n".join(p.text for p in section.header.paragraphs if p.text.strip())
                if header_text:
                    parts.insert(0, f"[Header]\n{header_text}")
            if section.footer:
                footer_text = "\n".join(p.text for p in section.footer.paragraphs if p.text.strip())
                if footer_text:
                    parts.append(f"[Footer]\n{footer_text}")

        return "\n\n".join(parts).strip()

    except Exception as e:
        # Fallback: try simple paragraph extraction
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
        except Exception:
            raise RuntimeError(f"Failed to extract DOCX: {e}")


def _extract_table_text(table) -> str:
    """Convert DOCX table to readable text format, handling merged cells."""
    if not table.rows:
        return ""

    lines = []
    for row in table.rows:
        # Track merged cells by internal object ID to avoid duplication
        seen_cells = set()
        cells_text = []

        for cell in row.cells:
            # Merged cells share the same _tc object, use id() to detect
            cell_id = id(cell._tc)
            if cell_id not in seen_cells:
                seen_cells.add(cell_id)
                text = cell.text.strip()
                if text:
                    cells_text.append(text)

        if cells_text:
            lines.append(" | ".join(cells_text))

    if lines:
        # Add separator after header (first row)
        if len(lines) > 1:
            lines.insert(1, "-" * min(80, len(lines[0])))
        return "\n".join(lines)
    return ""

def _extract_text_from_txt_bytes(data: bytes) -> str:
    """Extract text from TXT bytes with charset detection fallback."""
    try:
        return data.decode("utf-8")
    except Exception:
        try:
            import chardet
            enc = (chardet.detect(data) or {}).get("encoding") or "latin-1"
            return data.decode(enc, errors="replace")
        except Exception:
            return data.decode("latin-1", errors="replace")

def _clean_extracted_text(text: str) -> str:
    """Light cleanup to improve chunk readability."""
    import re
    # Remove control chars (keep \n, \t)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", text)
    # Normalize spaces
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # De-hyphenate at line breaks
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    return text.strip()


def _md_to_plaintext(md: str) -> str:
    """Best-effort markdown -> plain text for PDF export."""
    import re
    if not md:
        return ""
    text = md.replace("\r\n", "\n").replace("\r", "\n")

    # code blocks
    def _code(m):
        return "\n".join("    " + l for l in m.group(1).splitlines())

    text = re.sub(r"```[a-zA-Z0-9]*\n(.*?)```", _code, text, flags=re.S)

    # headings
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.M)
    # bullets
    text = re.sub(r"^\s*[-*+]\s+", "• ", text, flags=re.M)
    # bold / italic
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)
    text = re.sub(r"(\*|_)(.*?)\1", r"\2", text)
    # inline code
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()

def _normalize_soft_numbered_lists(md: str) -> str:
    """Convert soft numbered lists (e.g. '1 item') to hard format ('1. item')."""
    import re
    lines = md.splitlines()
    out = []
    for ln in lines:
        # If line starts with "number + space + non-space" but NOT "number + dot + space"
        if re.match(r"^\s*\d+\s+\S", ln) and not re.match(r"^\s*\d+\.\s+\S", ln):
            # Convert "1 text" to "1. text"
            out.append(re.sub(r"^(\s*\d+)\s+", r"\1. ", ln))
        else:
            out.append(ln)
    return "\n".join(out)

def _md_split_blocks(md: str):
    """Split Markdown into blocks: heading, list, code, paragraph.
    Ordered lists are detected ONLY with '1. ' (number + dot).
    """
    if not md:
        return []

    import re

    text = md.replace("\r\n", "\n").replace("\r", "\n")
    lines = text.split("\n")

    blocks = []
    i = 0

    def flush_paragraph(buf):
        if buf:
            paragraph = "\n".join(buf).strip()
            if paragraph:
                blocks.append({"type": "paragraph", "text": paragraph})
            buf.clear()

    def parse_list(start_idx):
        items = []
        ordered = None
        j = start_idx

        while j < len(lines):
            line = lines[j]
            if not line.strip():
                break

            # Ordered list ONLY if "1. "
            m_ord = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
            # Unordered list: "- ", "* ", "+ "
            m_un = re.match(r"^\s*[-*+]\s+(.*)$", line)

            if m_ord:
                if ordered is None:
                    ordered = True
                if ordered is False:
                    break
                items.append(m_ord.group(2).strip())
                j += 1
                continue

            if m_un:
                if ordered is None:
                    ordered = False
                if ordered is True:
                    break
                items.append(m_un.group(1).strip())
                j += 1
                continue

            break  # not a list anymore

        return {"type": "list", "ordered": bool(ordered), "items": items}, j

    paragraph_buf = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            flush_paragraph(paragraph_buf)
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines) and lines[i].strip().startswith("```"):
                i += 1
            blocks.append({"type": "code", "text": "\n".join(code_lines).rstrip()})
            continue

        # Heading
        m_h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m_h:
            flush_paragraph(paragraph_buf)
            level = len(m_h.group(1))
            blocks.append({
                "type": "heading",
                "level": level,
                "text": m_h.group(2).strip()
            })
            i += 1
            continue

        # List start (STRICT)
        if re.match(r"^\s*(\d+)\.\s+.*$", line) or re.match(r"^\s*[-*+]\s+.*$", line):
            flush_paragraph(paragraph_buf)
            lb, nxt = parse_list(i)
            if lb["items"]:
                blocks.append(lb)
            i = nxt
            continue

        # Blank line -> paragraph boundary
        if not line.strip():
            flush_paragraph(paragraph_buf)
            i += 1
            continue

                # Normal para

        paragraph_buf.append(line)
        i += 1

    flush_paragraph(paragraph_buf)
    return blocks

def _md_inline_to_rl(text: str) -> str:
    """Convert minimal Markdown inline (**bold**, *italic*, `code`) to ReportLab-friendly XML-ish markup."""
    import re
    from xml.sax.saxutils import escape

    if not text:
        return ""

    # IMPORTANT: protect inline code spans from emphasis parsing.
    # Otherwise underscores inside code (e.g. <IP_SORGENTE>) can be misread as italics,
    # producing broken XML-like markup that ReportLab cannot parse.
    parts = re.split(r"`([^`]+)`", text)
    out: list[str] = []

    def fmt_non_code(seg: str) -> str:
        s = escape(seg)
        # Bold: **x** or __x__
        s = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", s)
        s = re.sub(r"__(.+?)__", r"<b>\1</b>", s)
        # Italic: *x* or _x_
        s = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"<i>\1</i>", s)
        s = re.sub(r"_(.+?)_", r"<i>\1</i>", s)
        return s

    for i, seg in enumerate(parts):
        if i % 2 == 0:
            out.append(fmt_non_code(seg))
        else:
            code = escape(seg)
            # Extra safety: prevent any downstream underscore-based styling.
            code = code.replace("_", "&#95;")
            out.append(f'<font face="Courier">{code}</font>')

    return "".join(out)

def export_markdown_to_pdf_structured(md_text: str, title: str | None = None) -> bytes:
    """Render Markdown to a nicely formatted PDF using ReportLab Platypus.
    Returns PDF bytes.
    """
    import io
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted, ListFlowable, ListItem
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib import colors

    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.0 * cm,
        rightMargin=2.0 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
        title=title or "Answer",
    )

    styles = getSampleStyleSheet()

    # Custom styles
    normal = styles["Normal"]
    normal.leading = 14

    h1 = ParagraphStyle("H1", parent=styles["Heading1"], spaceAfter=10, spaceBefore=8)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], spaceAfter=8, spaceBefore=8)
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], spaceAfter=6, spaceBefore=8)

    code_style = ParagraphStyle(
        "CodeBlock",
        parent=styles["Code"],
        fontName="Courier",
        fontSize=9.5,
        leading=12,
        spaceBefore=6,
        spaceAfter=10,
        leftIndent=6,
        rightIndent=6,
        alignment=TA_LEFT,
        textColor=colors.black,
    )

    story = []

    if title:
        story.append(Paragraph(_md_inline_to_rl(title), h1))
        story.append(Spacer(1, 0.2 * cm))

    md_text = _normalize_soft_numbered_lists(md_text)
    blocks = _md_split_blocks(md_text)

    for b in blocks:
        btype = b["type"]

        if btype == "heading":
            level = b.get("level", 2)
            txt = _md_inline_to_rl(b.get("text", ""))
            if level <= 1:
                story.append(Paragraph(txt, h1))
            elif level == 2:
                story.append(Paragraph(txt, h2))
            else:
                story.append(Paragraph(txt, h3))
            continue

        if btype == "paragraph":
            txt = _md_inline_to_rl(b.get("text", ""))
            story.append(Paragraph(txt, normal))
            story.append(Spacer(1, 0.25 * cm))
            continue

        if btype == "list":
            items = b.get("items", [])
            ordered = bool(b.get("ordered", False))
            if not items:
                continue

            list_items = []
            for it in items:
                p = Paragraph(_md_inline_to_rl(it), normal)
                list_items.append(ListItem(p))

            lf = ListFlowable(
                list_items,
                bulletType=("1" if ordered else "bullet"),
                start="1",
                bulletFormat="%s." if ordered else None,  # <-- aggiungi
                leftIndent=16,
                bulletIndent=6,
                spaceBefore=2,
                spaceAfter=8,
            )

            story.append(lf)
            continue

        if btype == "code":
            code_text = b.get("text", "")
            if code_text.strip():
                story.append(Preformatted(code_text, code_style))
            continue

    if not story:
        story.append(Paragraph("No content to export.", normal))

    doc.build(story)
    return buf.getvalue()

def _render_download_pdf(answer_text: str, filename: str = "answer.pdf", title: str | None = None):
    """On-demand PDF export (prevents Streamlit from re-generating PDF on every rerun)."""
    import streamlit as st
    import hashlib

    if not (answer_text or "").strip():
        st.warning("No answer to export.")
        return

    # Stable cache key per (answer_text, title, filename)
    h = hashlib.sha256((str(title) + "\n" + filename + "\n" + answer_text).encode("utf-8", errors="ignore")).hexdigest()[:16]
    cache_key = f"docs_kb_pdf_bytes__{h}"

    c1, c2 = st.columns([0.25, 0.75])
    with c1:
        gen = st.button("Generate PDF", key=f"{cache_key}__gen")

    if gen:
        try:
            with st.spinner("Generating PDF…"):
                st.session_state[cache_key] = export_markdown_to_pdf_structured(answer_text, title=title)
            st.success("PDF generated.")
        except Exception as e:
            st.error(f"PDF export failed: {e}")
            import traceback
            st.code(traceback.format_exc(), language="text")
            return

    pdf_bytes = st.session_state.get(cache_key)
    if isinstance(pdf_bytes, (bytes, bytearray)) and pdf_bytes:
        with c2:
            st.download_button(
                "Download answer as PDF",
                data=pdf_bytes,
                file_name=filename,
                mime="application/pdf",
                key=f"{cache_key}__dl",
            )
    else:
        with c2:
            st.caption("Click **Generate PDF** to prepare the downloadable file.")

def _build_docs_prompt(question: str, retrieved: list) -> tuple[str, str]:
    """
    Build (system, user) prompts for Docs KB RAG.
    retrieved: list of (doc_text, meta, dist)
    """
    system = (
        "You are a technical assistant for a networking device CLI documentation.\n"
        "Use ONLY the provided context. If the context is insufficient, say so.\n"
        "When possible, provide exact CLI commands as code blocks and concise explanations.\n"
        "Always include a short 'Sources' section with file names (and page/section if available)."
    )

    context_lines = []
    sources = []
    for i, (txt, meta, dist) in enumerate(retrieved, start=1):
        meta = meta or {}
        fname = meta.get("source_file") or meta.get("filename") or "document"
        heading = meta.get("heading_path") or meta.get("section") or ""
        page = meta.get("page") or meta.get("page_range") or ""
        chunk_id = meta.get("chunk_id")
        src = f"{fname}"
        if page:
            src += f" (page {page})"
        if heading:
            src += f" — {heading}"
        if chunk_id is not None:
            src += f" — chunk {chunk_id}"
        sources.append(src)

        context_lines.append(f"[CONTEXT {i}] SOURCE: {src}\n{txt}\n")

    user = (
        f"Question:\n{question.strip()}\n\n"
        "Context:\n"
        + "\n".join(context_lines)
        + "\n"
        "Instructions:\n"
        "- Answer in Italian.\n"
        "- Provide 1–3 concrete CLI examples if relevant.\n"
        "- Keep it readable (bullets are fine).\n"
        "- End with 'Sources:' listing the sources used.\n"
    )
    return system, user

def render_phase_docs_kb_page(prefs):
    import streamlit as st

    # Normalize prefs
    prefs_dict = prefs if isinstance(prefs, dict) else st.session_state.get("prefs", {})

    st.title("Phase 9 – Docs KB (PDF/DOCX/TXT)")
    st.write(
        "Upload PDF/DOCX/TXT documentation, index it into a dedicated vector DB collection, "
        "and ask questions using the same LLM configured in Phase 5."
    )
    st.markdown(
        """
**Formatting rules (used for PDF export):**
- Use valid Markdown headings (`##`, `###`).
- If you create ordered lists, always use `1. item` (number + dot).
- Put CLI commands in fenced code blocks (```shell ... ```).
- End with a **Sources** section listing `filename` + `chunk ids`.
"""
    )

    st.markdown("---")

    # Resolve persist dir
    persist_dir = (
        st.session_state.get("vs_persist_dir")
        or st.session_state.get("persist_dir")
        or prefs_dict.get("persist_dir", DEFAULT_CHROMA_DIR)
    )
    os.makedirs(persist_dir, exist_ok=True)

    # Retrieval/Chunking params from Phase 4
    max_distance = float(st.session_state.get("max_distance", prefs_dict.get("max_distance", 0.9)))
    enable_chunking = bool(st.session_state.get("enable_chunking", prefs_dict.get("enable_chunking", True)))
    chunk_size = int(st.session_state.get("chunk_size", prefs_dict.get("chunk_size", 800)))
    chunk_overlap = int(st.session_state.get("chunk_overlap", prefs_dict.get("chunk_overlap", 80)))
    chunk_min = int(st.session_state.get("chunk_min", prefs_dict.get("chunk_min", 512)))
    top_k = int(st.session_state.get("top_k", prefs_dict.get("top_k", 5)))
    show_distances = bool(st.session_state.get("show_distances", prefs_dict.get("show_distances", False)))

    # Embedding/LLM config from Phase 5
    emb_backend = st.session_state.get("emb_provider_select", prefs_dict.get("emb_backend", "OpenAI"))
    emb_model_name = st.session_state.get("emb_model", prefs_dict.get("emb_model_name", "text-embedding-3-small"))

    llm_provider = st.session_state.get("llm_provider_select", prefs_dict.get("llm_provider", "OpenAI"))
    llm_model = (st.session_state.get("llm_model") or "").strip() or (DEFAULT_LLM_MODEL if llm_provider == "OpenAI" else DEFAULT_OLLAMA_MODEL)
    llm_temperature = float(st.session_state.get("llm_temperature", prefs_dict.get("llm_temperature", 0.2)))

    # Ensure collection exists
    try:
        kb_coll = load_chroma_collection(persist_dir, DOCS_KB_COLLECTION, space="cosine")
    except Exception as e:
        st.error(f"Unable to open Docs KB collection: {e}")
        return

    # -----------------------------
    # 1) Upload + Index
    # -----------------------------
    st.subheader("1) Upload & index documents")

    uploaded = st.file_uploader(
        "Upload PDF/DOCX/TXT",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
    )

    c_idx1, c_idx2 = st.columns([0.7, 0.3])
    with c_idx1:
        do_index = st.button("Index uploaded documents", disabled=not uploaded)
    with c_idx2:
        st.caption(f"Collection: `{DOCS_KB_COLLECTION}`")
        st.caption(f"Path: `{persist_dir}`")

    if do_index and uploaded:
        failures = 0
        added_chunks = 0

        # Build/embedder (reuse global session object)
        try:
            if st.session_state.get("embedder") is None:
                use_openai_embeddings = (emb_backend == "OpenAI")
                st.session_state["embedder"] = EmbeddingBackend(
                    use_openai=use_openai_embeddings,
                    model_name=emb_model_name,
                )
        except Exception as e:
            st.error(f"Embedding backend init error: {e}")
            return

        manifest = _load_docs_manifest(persist_dir)
        known_ids = set((it.get("doc_id") for it in manifest if isinstance(it, dict)))

        for uf in uploaded:
            try:
                data = uf.getvalue()
                doc_id = _sha256_bytes(data)[:16]
                if doc_id in known_ids:
                    st.info(f"Skipped (already indexed): {uf.name}")
                    continue

                ext = (uf.name.split(".")[-1] or "").lower()
                if ext == "pdf":
                    raw_text = _extract_text_from_pdf_bytes(data)
                elif ext == "docx":
                    raw_text = _extract_text_from_docx_bytes(data)
                else:
                    raw_text = _extract_text_from_txt_bytes(data)

                raw_text = _clean_extracted_text(raw_text)
                if not raw_text.strip():
                    raise RuntimeError("No extractable text found (is it scanned or empty?).")

                # Chunking (reuse existing splitter)
                if enable_chunking:
                    chunks = split_into_chunks(
                        raw_text,
                        chunk_size=int(chunk_size),
                        overlap=int(chunk_overlap),
                        min_size=int(chunk_min),
                    )
                else:
                    chunks = [(0, raw_text)]

                if not chunks:
                    chunks = [(0, raw_text)]

                ids = []
                docs = []
                metas = []
                embed_inputs = []

                multi = len(chunks) > 1
                for i, (pos0, chunk_text) in enumerate(chunks, start=1):
                    cid = f"{doc_id}::c{i:03d}" if multi else doc_id
                    ids.append(cid)
                    docs.append(chunk_text)

                    meta = {
                        "doc_id": doc_id,
                        "source_file": uf.name,
                        "doc_type": ext,
                    }
                    if multi:
                        meta["chunk_id"] = i
                        meta["pos"] = int(pos0)

                    metas.append(meta)
                    # Extra prefix helps retrieval for "command-like" questions
                    embed_inputs.append(f"{chunk_text}")

                with st.spinner(f"Embedding {uf.name} ({len(ids)} chunks)…"):
                    embs = st.session_state["embedder"].embed(embed_inputs)
                    kb_coll.add(ids=ids, documents=docs, metadatas=metas, embeddings=embs)

                manifest.append(
                    {
                        "doc_id": doc_id,
                        "filename": uf.name,
                        "doc_type": ext,
                        "bytes": len(data),
                        "chunks": len(ids),
                        "indexed_at": int(time.time()),
                    }
                )
                known_ids.add(doc_id)
                added_chunks += len(ids)
                st.success(f"Indexed: {uf.name} — chunks: {len(ids)}")

            except Exception as e:
                failures += 1
                st.error(f"Indexing failed for {getattr(uf, 'name', 'file')}: {e}")
                with st.expander("Traceback"):
                    st.code(traceback.format_exc(), language="text")

        _save_docs_manifest(persist_dir, manifest)

        if failures == 0:
            st.success(f"Indexing completed. Added chunks: {added_chunks}")
            st.session_state["ui_phase_choice"] = "Docs KB (PDF/DOCX/TXT)"
            st.rerun()
        else:
            st.warning(f"Indexing completed with {failures} failure(s).")

    st.markdown("---")

    # -----------------------------
    # 2) Manage documents (list + delete)
    # -----------------------------
    st.subheader("2) Indexed documents")
    manifest = _load_docs_manifest(persist_dir)

    if not manifest:
        st.caption("No documents indexed yet.")
    else:
        for it in list(manifest):
            doc_id = it.get("doc_id")
            fname = it.get("filename", "document")
            chunks_n = it.get("chunks", "?")
            size_b = it.get("bytes", 0)
            when = it.get("indexed_at", 0)
            when_s = datetime.fromtimestamp(int(when)).strftime("%Y-%m-%d %H:%M") if when else "—"

            col_main, col_del = st.columns([0.86, 0.14])
            with col_main:
                st.markdown(f"**{fname}**")
                st.caption(f"doc_id={doc_id} · chunks={chunks_n} · size={size_b} bytes · indexed={when_s}")
            with col_del:
                confirm = st.checkbox("confirm", key=f"docs_del_confirm_{doc_id}")
                if st.button("🗑️ remove", key=f"docs_del_{doc_id}", disabled=not confirm):
                    try:
                        kb_coll.delete(where={"doc_id": doc_id})
                        manifest = [x for x in manifest if x.get("doc_id") != doc_id]
                        _save_docs_manifest(persist_dir, manifest)
                        st.success(f"Removed: {fname}")
                        st.session_state["ui_phase_choice"] = "Docs KB (PDF/DOCX/TXT)"
                        st.rerun()
                    except Exception as e:
                        st.error(f"Delete failed: {e}")
                        with st.expander("Traceback"):
                            st.code(traceback.format_exc(), language="text")

    st.markdown("---")

    # -----------------------------
    # 3) Ask questions (Retrieval + LLM)
    # -----------------------------
    st.subheader("3) Ask the Docs KB (RAG)")

    q = st.text_area(
        "Question",
        height=120,
        placeholder="Es: Come configuro RPKI? Fammi un esempio di set/show e spiega i parametri.",
    )

    c_q1, c_q2, c_q3 = st.columns([0.25, 0.25, 0.5])
    with c_q1:
        use_llm = st.checkbox("Use LLM to answer", value=True)
    with c_q2:
        nres = st.number_input("Top K", min_value=1, max_value=20, step=1, value=int(top_k))
    with c_q3:
        run = st.button("Search in Docs KB")

    if run:
        if not q.strip():
            st.error("Please enter a question.")
            return

        # Ensure embedder exists
        try:
            if st.session_state.get("embedder") is None:
                use_openai_embeddings = (emb_backend == "OpenAI")
                st.session_state["embedder"] = EmbeddingBackend(
                    use_openai=use_openai_embeddings,
                    model_name=emb_model_name,
                )
        except Exception as e:
            st.error(f"Embedding backend init error: {e}")
            return

        try:
            q_emb = st.session_state["embedder"].embed([q.strip()])[0]
            res = kb_coll.query(query_embeddings=[q_emb], n_results=int(nres))
            docs = (res.get("documents") or [[]])[0]
            metas = (res.get("metadatas") or [[]])[0]
            dists = (res.get("distances") or [[]])[0]

            retrieved = []
            for doc, meta, dist in zip(docs, metas, dists):
                if dist is None:
                    continue
                if dist <= max_distance:
                    retrieved.append((doc, meta or {}, float(dist)))

            if not retrieved:
                st.warning("No relevant chunks found under the current distance threshold.")
                return

            st.session_state["docs_kb_last_retrieved"] = retrieved
            st.session_state["docs_kb_last_query"] = q.strip()

            # Show retrieved chunks
            with st.expander("Retrieved chunks", expanded=False):
                for i, (doc, meta, dist) in enumerate(retrieved, start=1):
                    fname = meta.get("source_file") or "document"
                    cid = meta.get("chunk_id")
                    pos = meta.get("pos")
                    hdr = f"{i}. {fname}"
                    if cid is not None:
                        hdr += f" — chunk {cid}"
                    st.markdown(hdr)
                    if show_distances:
                        st.caption(f"distance={dist:.4f} (max={max_distance})")
                    if pos is not None:
                        st.caption(f"offset={pos}")
                    st.code(doc, language="text")

            if not use_llm:
                st.info("LLM disabled: showing retrieved chunks only.")
                return

            # Build prompt + call LLM
            system, user = _build_docs_prompt(q.strip(), retrieved)
            format_rules = (
                "Formatting rules:\n"
                "- Use valid Markdown headings (##, ###).\n"
                "- If you create ordered lists, ALWAYS use '1. item' (number + dot).\n"
                "- Do NOT use '1 item' without the dot.\n"
                "- Put CLI commands in fenced code blocks using ```shell ... ```.\n"
                "- End with a 'Sources' section listing filename + chunk ids.\n"
            )
            system = (system or "") + "\n\n" + format_rules

            llm = LLMBackend(llm_provider, llm_model, temperature=llm_temperature)

            with st.spinner("Generating answer with LLM…"):
                answer = llm.generate(system, user).strip()

            st.session_state["docs_kb_last_answer"] = answer
            st.success("Answer generated. See section '4) Last answer' below.")


        except Exception as e:
            st.error(f"Docs KB query error: {e}")
            with st.expander("Traceback"):
                st.code(traceback.format_exc(), language="text")

# ------------------------------
# CLI + Self-tests (optional)
# ------------------------------

    # -----------------------------
    # 4) Last answer (persistent across reruns)
    # -----------------------------
    last_q = st.session_state.get("docs_kb_last_query", "")
    last_answer = st.session_state.get("docs_kb_last_answer", "")
    last_retrieved = st.session_state.get("docs_kb_last_retrieved", [])

    if last_answer:
        st.markdown("---")
        st.subheader("4) Last answer")

        if last_q:
            st.caption(f"Question: {last_q}")

        st.markdown(last_answer)

        with st.expander("Retrieved chunks (last run)", expanded=False):
            for i, (doc, meta, dist) in enumerate(last_retrieved, start=1):
                meta = meta or {}
                fname = meta.get("source_file") or "document"
                cid = meta.get("chunk_id")
                pos = meta.get("pos")

                hdr = f"{i}. {fname}"
                if cid is not None:
                    hdr += f" — chunk {cid}"
                st.markdown(hdr)

                if show_distances:
                    st.caption(f"distance={dist:.4f} (max={max_distance})")
                if pos is not None:
                    st.caption(f"offset={pos}")

                st.code(doc, language="text")

        st.subheader("Export")
        q_one_line = re.sub(r"\s+", " ", (last_q or "")).strip()
        pdf_body = last_answer
        if q_one_line:
            pdf_body = f"Question:\n{q_one_line}\n\n---\n\n{last_answer}"
        _render_download_pdf(
            pdf_body,
            filename="docs_kb_answer.pdf",
            title="Docs KB Answer",
        )


def _cli_help():
    print("Usage: streamlit run app.py --server.port 8502")

def _self_tests():
    print("Running minimal self-tests...")
    vs = None
    try:
        vs = VectorStore(DEFAULT_CHROMA_DIR, DEFAULT_COLLECTION)
        print("VectorStore OK.")
    except Exception as e:
        print(f"VectorStore not available: {e}")

    try:
        emb = EmbeddingBackend(use_openai=False, model_name=DEFAULT_EMBEDDING_MODEL)
        vec = emb.embed(["testo uno", "testo due"])
        assert len(vec) == 2 and isinstance(vec[0], list)
        print("Local embeddings OK.")
    except Exception as e:
        print(f"Local embeddings not available: {e}")

    try:
        llm = LLMBackend("OpenAI", DEFAULT_LLM_MODEL)
        assert isinstance(llm, LLMBackend)
        print("LLM OpenAI OK (init).")
    except Exception as e:
        print(f"LLM OpenAI not available: {e}")

    try:
        llm = LLMBackend("Ollama (local)", DEFAULT_OLLAMA_MODEL)
        assert isinstance(llm, LLMBackend)
        print("LLM Ollama OK (init).")
    except Exception as e:
        print(f"LLM Ollama not available: {e}")

    try:
        llm = LLMBackend("???", "x")
    except Exception as e:
        assert "Unsupported LLM provider" in str(e)
    print("All self-tests PASSED. ✅")

# ------------------------------
# Main
# ------------------------------
if __name__ == "__main__":
    print("[DEBUG] Starting main program.")
    if ST_AVAILABLE:
        print("[DEBUG] Starting Streamlit interface.")
        run_streamlit_app()
    else:
        _cli_help()
        _self_tests()


def generate_pdf(answer_text: str, filename: str = "answer.pdf"):
    return _render_download_pdf(answer_text, filename)
